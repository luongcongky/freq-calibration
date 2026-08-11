"""
unit_test/test_report_templates_generic.py
==============================================
Test core/report_templates/generic.py::GenericReportTemplate +
core/report_templates/__init__.py's fallback đăng ký — 1 biểu mẫu hoàn
toàn data-driven (không có class Python riêng) phải hoạt động đúng qua
interface BaseReportTemplate (REGISTRY hiện rỗng — mọi template, kể cả
TEMPLATE_FREQ/TEMPLATE_POWER, đều đi qua cơ chế fallback này).
"""

import json
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from core import table_engine
from core.table_descriptor import TableDescriptor, RowDef, ColumnDef
from core import table_wizard_io as wio
from core.session import CalibrationSession, SessionMeta, DUTInfo


def _build_fixture_template(base_dir: Path, template_id: str):
    """Dựng 1 biểu mẫu tối giản: meta.json + tables/T1.json + bienban.docx
    + gcnkd.docx (cùng khuôn 1 bảng T1)."""
    tpl_dir = base_dir / template_id
    (tpl_dir / "tables").mkdir(parents=True)

    meta = {
        "template_id": template_id, "template_name": "Mẫu thử nghiệm",
        "kind": "kiem_dinh", "dut_models": ["X1"], "standard": "TEST : 2026",
        "measurement_range": "0-1 GHz",
    }
    (tpl_dir / "meta.json").write_text(json.dumps(meta, ensure_ascii=False), encoding="utf-8")

    descriptor = TableDescriptor(
        schema_version=1, table_id="T1", name="Bảng thử nghiệm", order=1,
        scenario_file="", layout="repeated_rows", value_unit="Hz",
        rows=[RowDef(key="f1", freq_set=1e6, reference=1e6, raw_count=1)],
        columns=[
            ColumnDef(role="freq_set", format="freq", jinja_field="freq_str", scope="row", col=0),
            ColumnDef(role="measured", format="hz_measured", jinja_field="value_str", scope="row", col=1),
        ],
        pass_rule={"type": "none"}, merge=[],
    )
    wio.write_descriptor_json(descriptor, tpl_dir / "tables")

    def _build_docx(path):
        doc = Document()
        doc.add_paragraph("{% if tables.T1.enabled %}")
        tbl = doc.add_table(rows=4, cols=2)
        tbl.cell(0, 0).text = "Tần số"
        tbl.cell(0, 1).text = "Giá trị đo"
        tbl.cell(1, 0).text = "{%tr for row in tables.T1.rows %}"
        tbl.cell(2, 0).text = "{{ row.freq_str }}"
        tbl.cell(2, 1).text = "{{ row.value_str }}"
        tbl.cell(3, 0).text = "{%tr endfor %}"
        table_engine.mark_table(tbl, "T1")
        doc.add_paragraph("{% endif %}")
        sig = doc.add_table(rows=1, cols=1)
        sig.cell(0, 0).text = "Chữ ký"
        doc.save(str(path))

    _build_docx(tpl_dir / "bienban.docx")
    _build_docx(tpl_dir / "gcnkd.docx")
    return tpl_dir


@pytest.fixture
def fixture_registry(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)
    _build_fixture_template(tmp_path, "TEST_GENERIC_TPL")
    return tmp_path


def test_load_generic_meta_and_discover(fixture_registry):
    from core.report_templates.generic import load_generic_meta, discover_generic_template_ids
    meta = load_generic_meta("TEST_GENERIC_TPL")
    assert meta["template_name"] == "Mẫu thử nghiệm"
    assert discover_generic_template_ids() == ["TEST_GENERIC_TPL"]
    assert load_generic_meta("KHONG_TON_TAI") is None


def test_get_template_fallback_and_list_templates(fixture_registry):
    from core.report_templates import get_template, list_templates
    from core.report_templates.generic import GenericReportTemplate

    tpl = get_template("TEST_GENERIC_TPL")
    assert isinstance(tpl, GenericReportTemplate)
    assert tpl.TEMPLATE_NAME == "Mẫu thử nghiệm"
    assert tpl.DUT_MODELS == ["X1"]

    ids = [tid for tid, _ in list_templates()]
    assert "TEST_GENERIC_TPL" in ids

    with pytest.raises(KeyError):
        get_template("KHONG_TON_TAI_LUON")


def test_generic_template_full_roundtrip(fixture_registry, tmp_path):
    from core.report_templates import get_template

    tpl = get_template("TEST_GENERIC_TPL")
    tests = tpl.default_tests()
    assert len(tests) == 1 and tests[0].table_id == "T1"

    session = CalibrationSession(
        template_id="TEST_GENERIC_TPL",
        meta=SessionMeta(dut=DUTInfo(serial="SN1"), operator="A", reviewer="B",
                          date=date(2026, 7, 30)),
        tests=tests,
    )
    tpl.fill_session_defaults(session)
    assert session.meta.dut.model == "X1"
    assert session.meta.dut.measurement_range == "0-1 GHz"

    from core.scenario_runner import StepResult
    session.tests[0].step_results = [StepResult(action="report_val", value=1000000.5, ok=True)]
    session.tests[0].result_table = tpl.map_test_result(session.tests[0])
    for r in session.tests[0].result_table.rows:
        r.confirmed = True

    out_bb = tpl.generate_bienban(session, tmp_path / "out_bienban.docx")
    out_gcn = tpl.generate_gcnkd(session, tmp_path / "out_gcnkd.docx")

    for out in (out_bb, out_gcn):
        doc = Document(str(out))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text
        assert "{{" not in full_text and "{%" not in full_text
