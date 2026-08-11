"""
unit_test/test_generic_seed_templates.py
===========================================
Test 2 template mới dựng qua scripts/build_generic_seed_templates.py
("TEMPLATE_FREQ" — đo tần số, thay QTKD_2461_CNT90XL; "TEMPLATE_POWER" —
hiệu chuẩn công suất, thay QTHC_2515_NRP2) — cả 2 đều GenericReportTemplate
data-driven, không còn class Python riêng.

Bỏ qua toàn bộ nếu templates/<id>/bienban.docx chưa được sinh (chạy
`python scripts/build_generic_seed_templates.py` trước).
"""

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from core.report_templates import get_template
from core.report_templates.generic import GenericReportTemplate
from core.scenario_runner import StepResult
from core.session import CalibrationSession, DUTInfo, SessionMeta

FREQ_ID = "TEMPLATE_FREQ"
POWER_ID = "TEMPLATE_POWER"
FREQ_BIENBAN = Path("templates") / FREQ_ID / "bienban.docx"
POWER_BIENBAN = Path("templates") / POWER_ID / "bienban.docx"

pytestmark = pytest.mark.skipif(
    not (FREQ_BIENBAN.exists() and POWER_BIENBAN.exists()),
    reason="Chưa sinh templates/ — chạy scripts/build_generic_seed_templates.py trước")


def _no_leftover_tags(doc: Document) -> bool:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    return "{{" not in full_text and "{%" not in full_text


def _meta():
    return SessionMeta(
        dut=DUTInfo(name="Thiết bị thử", model="MODEL-X", serial="SN1", manufacturer="ACME",
                    manufacture_year="2024", owner="Công ty ABC", measurement_range="X-Y"),
        operator="Nguyễn Văn A", reviewer="Trần Thị B", manager="Lê Văn C",
        conclusion="Đạt yêu cầu kỹ thuật đo lường", date=date(2026, 8, 5),
    )


def _run(tpl, test, values):
    test.step_results = [StepResult(action="report_val", value=v, ok=True) for v in values]
    test.result_table = tpl.map_test_result(test)
    for r in test.result_table.rows:
        r.confirmed = True
    return test


# ---------------------------------------------------------------------------
# TEMPLATE_FREQ
# ---------------------------------------------------------------------------

def test_freq_registered_with_8_tables_in_order():
    tpl = get_template(FREQ_ID)
    assert isinstance(tpl, GenericReportTemplate)
    assert [t.table_id for t in tpl.default_tests()] == \
        ["A1", "A2", "A3", "A4", "A5", "A6", "A7", "A8"]


def test_freq_bienban_renders_clean_and_merges_a2(tmp_path):
    from core import table_engine

    tpl = get_template(FREQ_ID)
    tests = {t.table_id: t for t in tpl.default_tests()}
    # A1: measured_count=5 (fCi) + 2 field kịch bản TỰ TÍNH đẩy thêm (fC rồi
    # δf) — raw_count=7, xem FREQ_COMPUTED_ERROR_EXTRA_FORMATS trong
    # scripts/build_generic_seed_templates.py.
    fci = [10_000_000.0 + i for i in range(5)]
    fc, delta_f = 10_000_000.2, 2e-8
    _run(tpl, tests["A1"], fci + [fc, delta_f])
    _run(tpl, tests["A2"], [12.0] * 12)                             # 12 dòng x raw_count=1

    session = CalibrationSession(template_id=FREQ_ID, meta=_meta())
    session.tests = list(tests.values())
    for t in session.tests:
        t.enabled = t.table_id in ("A1", "A2")

    out = tpl.generate_bienban(session, tmp_path / "freq_bienban.docx")
    doc = Document(str(out))
    assert _no_leftover_tags(doc)

    # Không còn marker ẩn (gộp ô giờ làm THẲNG trong Word lúc dựng file, xem
    # scripts/build_generic_seed_templates.py) — tìm bảng theo VỊ TRÍ: bảng
    # 0 = header nhà nước, sau đó 1 bảng/mục ĐANG BẬT theo đúng thứ tự A1..A8.
    a1 = doc.tables[1]
    a2 = doc.tables[2]
    assert len(a1.rows) == 2 + 5   # header + 5 lần đo fCi + 1 dòng fC/δf riêng
    # 5 cột: Tần số thiết lập | fCi | fC | δf | Sai số cho phép — khớp mẫu
    # giấy gốc QTKĐ 2.461:2018 Phụ lục A (Bảng A1).
    assert [c.text for c in a1.rows[0].cells] == [
        "Tần số thiết lập", "Tần số đo được\ntrên CNT-90XL\n(fCi)",
        "Tần số đo được\ntrên CNT-90XL\n(fC)", "Sai số tần số\n(δf)",
        "Sai số\ncho phép\n(δfcp)"]
    assert a1.cell(1, 1).text == table_engine._format("hz_measured", fci[0])
    assert a1.cell(2, 1).text == table_engine._format("hz_measured", fci[1])
    # fC/δf: dòng RIÊNG (dòng 6) sau khối 5 dòng fCi — KHÔNG gộp dọc xuyên
    # suốt khối fCi (docxtpl đọc tag ở dòng ĐẦU vùng gộp, sẽ xen giữa các
    # fCi nếu gộp — xem chú thích trong scripts/build_generic_seed_templates.py).
    assert a1.cell(6, 2).text == table_engine._format("hz_measured", fc)
    assert a1.cell(6, 3).text == table_engine._format("sci", delta_f)
    assert a1.cell(6, 4).text == "± 2,4×10⁻⁷"
    # Cột "Tần số thiết lập" vẫn gộp xuyên suốt cả khối (5 dòng fCi + dòng
    # fC/δf) vì đây là TĨNH (không phải tag), không bị vấn đề thứ tự trên.
    assert a1.cell(1, 0)._tc is a1.cell(6, 0)._tc

    assert len(a2.rows) == 13   # header + 12 dòng
    assert a2.cell(0, 1).text == "Độ nhạy đo được"
    assert a2.cell(1, 1).text == "12 mVrms"   # report_val() đã format đúng đơn vị
    # Cột giới hạn (col 2) gộp theo nhóm liên tiếp cùng limit: 9 dòng đầu
    # "<= 15 mVrms" (100kHz..150MHz) rồi 3 dòng "<= 25 mVrms" (200/250/300MHz)
    same_group = a2.cell(1, 2)._tc is a2.cell(2, 2)._tc
    assert same_group, "2 dòng đầu (cùng giới hạn 15mVrms) phải được gộp ô"


def test_freq_bienban_a5_uses_two_report_val_per_row_for_value_and_error(tmp_path):
    """Bảng A5 (mẫu giấy gốc có cột 'Sai số đo tần số' cạnh giá trị đo) —
    kịch bản tự tính sai số rồi đẩy THÊM 1 report_val()/dòng (measured_count=1,
    raw_count=2) — Biên Bản vẫn chỉ echo lại đúng report_val(), không tự
    tính gì; value_measured/passed chỉ lấy từ slot ĐẦU (giá trị đo), không
    trộn với slot sai số."""
    from core import table_engine

    tpl = get_template(FREQ_ID)
    tests = {t.table_id: t for t in tpl.default_tests()}
    a5 = tests["A5"]
    # 11 dòng x 2 report_val (đo được rồi sai số kịch bản tự tính).
    values = []
    for i in range(11):
        values += [5.0 + i * 1e-8, 1.2e-8]   # (đo được, sai số) xen kẽ
    _run(tpl, a5, values)

    session = CalibrationSession(template_id=FREQ_ID, meta=_meta())
    session.tests = list(tests.values())
    for t in session.tests:
        t.enabled = t.table_id == "A5"

    out = tpl.generate_bienban(session, tmp_path / "freq_bienban_a5.docx")
    doc = Document(str(out))
    assert _no_leftover_tags(doc)

    a5_tbl = doc.tables[1]
    assert [c.text for c in a5_tbl.rows[0].cells] == [
        "Tần số thiết lập", "Tần số đo được trên kênh A (fđo)",
        "Sai số đo tần số\n(δf)", "Sai số\ncho phép"]
    assert a5_tbl.cell(1, 1).text == table_engine._format("hz_measured", 5.0)
    assert a5_tbl.cell(1, 2).text == table_engine._format("sci", 1.2e-8)
    assert a5_tbl.cell(1, 3).text == "± 2,4×10⁻⁷"

    # value_measured/passed dùng ĐÚNG slot "đo được", không bị slot sai số
    # (đã đẩy kèm) làm lệch trung bình.
    row0 = a5.result_table.rows[0]
    assert row0.value_measured == 5.0
    assert row0.passed is True


def test_freq_gcnkd_uses_result_cursor_and_hides_disabled_rows(tmp_path):
    tpl = get_template(FREQ_ID)
    tests = {t.table_id: t for t in tpl.default_tests()}
    _run(tpl, tests["A5"], [5.0000012] + [x for x in range(10)])  # đủ 11 dòng raw_count=1
    session = CalibrationSession(template_id=FREQ_ID, meta=_meta())
    session.tests = list(tests.values())
    for t in session.tests:
        t.enabled = t.table_id == "A5"

    out = tpl.generate_gcnkd(session, tmp_path / "freq_gcnkd.docx")
    doc = Document(str(out))
    assert _no_leftover_tags(doc)

    summary_rows = [[c.text for c in r.cells] for r in doc.tables[1].rows]
    # Chỉ đúng 1 dòng dữ liệu (A5) + 1 dòng header — mọi bảng khác đang tắt phải ẩn hẳn
    assert len(summary_rows) == 2, summary_rows
    assert summary_rows[1][0] == "5.Xác định sai số đo tần số kênh A"
    assert summary_rows[1][1] in ("Đạt", "Không đạt")
    assert summary_rows[1][2] == "± 2,4×10⁻⁷"


def test_freq_all_tables_disabled_gcnkd_has_only_header(tmp_path):
    tpl = get_template(FREQ_ID)
    session = CalibrationSession(template_id=FREQ_ID, meta=_meta())
    session.tests = tpl.default_tests()
    for t in session.tests:
        t.enabled = False

    out = tpl.generate_gcnkd(session, tmp_path / "freq_gcnkd_empty.docx")
    doc = Document(str(out))
    assert _no_leftover_tags(doc)
    summary_rows = [[c.text for c in r.cells] for r in doc.tables[1].rows]
    assert len(summary_rows) == 1   # chỉ còn dòng tiêu đề


# ---------------------------------------------------------------------------
# TEMPLATE_POWER
# ---------------------------------------------------------------------------

def test_power_registered_with_3_tables_in_order():
    tpl = get_template(POWER_ID)
    assert isinstance(tpl, GenericReportTemplate)
    assert [t.table_id for t in tpl.default_tests()] == ["A1", "A2", "A3"]


def test_power_bienban_renders_clean_and_merges_a3_freq_group(tmp_path):
    tpl = get_template(POWER_ID)
    tests = {t.table_id: t for t in tpl.default_tests()}
    # A1: measured_count=10 (2 nhóm 5) + 2 field kịch bản TỰ TÍNH đẩy thêm
    # (TB rồi Độ KĐBĐ) — raw_count=12, xem POWER_COMPUTED_EXTRA_FORMATS
    # trong scripts/build_generic_seed_templates.py.
    a1_measured = [0.00098, 0.00099, 0.00100, 0.00099, 0.00098,
                   0.00099, 0.00100, 0.00099, 0.00098, 0.00099]
    _run(tpl, tests["A1"], a1_measured + [0.00099, 0.00002])
    # A3: mỗi dòng 5 lần đo (measured_count=5) + 2 field kịch bản tự tính -> 7/dòng.
    a3_values = []
    for _ in range(48):
        a3_values += [-30.0] * 5 + [-30.0, 0.5]
    _run(tpl, tests["A3"], a3_values)

    session = CalibrationSession(template_id=POWER_ID, meta=_meta())
    session.tests = list(tests.values())
    for t in session.tests:
        t.enabled = t.table_id in ("A1", "A3")

    out = tpl.generate_bienban(session, tmp_path / "power_bienban.docx")
    doc = Document(str(out))
    assert _no_leftover_tags(doc)

    # Bảng 0 = header nhà nước, sau đó 1 bảng/mục ĐANG BẬT theo thứ tự A1..A3
    # (A2 đang tắt trong test này -> chỉ còn A1 rồi A3).
    a1 = doc.tables[1]
    a3 = doc.tables[2]
    assert len(a1.rows) == 4   # header + 2 nhóm 5 lần đo + 1 dòng Trung Bình
    assert a1.cell(3, 0).text == "Trung Bình"
    assert len(a3.rows) == 49   # header + 48 dòng
    # 6 dòng đầu (cùng tần số 50MHz, 6 mức dBm) phải gộp cột 0
    same_group = a3.cell(1, 0)._tc is a3.cell(2, 0)._tc
    assert same_group, "Các dòng cùng tần số phải được gộp cột 'Tần số thiết lập'"


def test_power_gcnkd_matches_official_blank_form_no_result_table(tmp_path):
    """Mẫu Phụ lục B (QTHC 2.515:2021) CHỈ có đúng 1 dòng "Kết quả (Results):"
    — KHÔNG có bảng A1/A2/A3 chi tiết nào (khác Biên Bản). Bảng TB/Số hiệu
    chỉnh/Độ KĐBĐ kiểu cũ đã bị bỏ theo đúng yêu cầu khách hàng khớp hệt ảnh
    mẫu — render sạch dù KHÔNG tham chiếu tables.X.* nào. "Kết quả (Results):"
    PHẢI gắn header.conclusion (ô "Kết luận" tự động ở Bước 3) — để trống
    hoàn toàn sẽ khiến GCN xuất ra không có kết luận cuối cùng nào (bug đã
    bị khách hàng phát hiện khi rà soát 1 phiên thật)."""
    tpl = get_template(POWER_ID)
    tests = {t.table_id: t for t in tpl.default_tests()}
    values = []
    for _ in range(13):
        values += [1.0] * 5 + [1.0, 0.05]
    _run(tpl, tests["A2"], values)
    meta = _meta()
    meta.conclusion = "Không đạt yêu cầu kỹ thuật đo lường"
    session = CalibrationSession(template_id=POWER_ID, meta=meta)
    session.tests = list(tests.values())
    for t in session.tests:
        t.enabled = t.table_id == "A2"

    out = tpl.generate_gcnkd(session, tmp_path / "power_gcnkd.docx")
    doc = Document(str(out))
    assert _no_leftover_tags(doc)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Kết quả (Results): Không đạt yêu cầu kỹ thuật đo lường" in full_text
    assert len(doc.tables) == 2   # chỉ 2 bảng chữ ký + footer, không có bảng kết quả nào
    assert "Trang: 01/03" in doc.tables[1].cell(0, 0).text


def test_power_bienban_uses_report_val_not_gcn_cursors(tmp_path):
    """Biên Bản phải render sạch dù KHÔNG có dữ liệu gcn_avg/gcn_error nào
    được xác nhận riêng — vì Biên Bản không tham chiếu 2 field đó chút nào,
    chỉ dùng report_val() (đúng những gì kịch bản đã đẩy)."""
    tpl = get_template(POWER_ID)
    tests = {t.table_id: t for t in tpl.default_tests()}
    values = []
    for _ in range(13):
        values += [2.0] * 5 + [2.0, 0.05]
    _run(tpl, tests["A2"], values)
    session = CalibrationSession(template_id=POWER_ID, meta=_meta())
    session.tests = list(tests.values())
    for t in session.tests:
        t.enabled = t.table_id == "A2"

    out = tpl.generate_bienban(session, tmp_path / "power_bienban_2.docx")
    doc = Document(str(out))
    assert _no_leftover_tags(doc)
    a2 = doc.tables[1]   # bảng 0 = header nhà nước
    from core import table_engine
    # Ô lần đo đầu tiên phải đúng giá trị thô đã đẩy (2.0), không phải TB.
    assert a2.cell(1, 1).text == table_engine._format("dbm", 2.0)
