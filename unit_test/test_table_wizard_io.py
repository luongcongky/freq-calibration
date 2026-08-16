"""
unit_test/test_table_wizard_io.py
====================================
Test phần "không Qt" của "Thêm bảng báo cáo mới" (core/table_wizard_io.py)
— dựng TableDescriptor từ dữ liệu người dùng nhập ở màn hình review, ghi
JSON, đọc file .docx đã gắn tag tay để gợi ý điền form/cảnh báo mã bảng
thiếu. KHÔNG còn quét/đoán cấu trúc bảng hay tự động chèn tag (quản trị
viên tự gõ tag Jinja trực tiếp trong Word).
"""

from pathlib import Path

import pytest
from docx import Document

from core import table_wizard_io as wio
from core.table_descriptor import load_table_descriptor


@pytest.mark.parametrize("text,expected", [
    ("100000", 100000.0), ("1,5", 1.5), (" 42 ", 42.0),
    ("100 kHz", None), ("≤ 15 mVrms", None), ("", None), (None, None),
])
def test_guess_bare_number(text, expected):
    assert wio.guess_bare_number(text) == expected


# ---------------------------------------------------------------------------
# Validate
# ---------------------------------------------------------------------------

def test_validate_table_id_available(tmp_path):
    assert wio.validate_table_id_available(tmp_path, "") is not None
    assert wio.validate_table_id_available(tmp_path, "A 9") is not None
    assert wio.validate_table_id_available(tmp_path, "A9") is None
    (tmp_path / "A9.json").write_text("{}", encoding="utf-8")
    assert wio.validate_table_id_available(tmp_path, "A9") is not None


def test_validate_rows():
    pass_rule = {"type": "relative_error_vs_fixed_limit", "params": {"fixed_limit": 1e-7, "limit_str": "x"}}
    rows_missing_ref = [wio.WizardRowSpec(key="a", reference=None)]
    assert wio.validate_rows(rows_missing_ref, pass_rule) is not None
    rows_ok = [wio.WizardRowSpec(key="a", reference=1e6)]
    assert wio.validate_rows(rows_ok, pass_rule) is None
    rows_dup_key = [wio.WizardRowSpec(key="a", reference=1e6), wio.WizardRowSpec(key="a", reference=2e6)]
    assert wio.validate_rows(rows_dup_key, pass_rule) is not None


# ---------------------------------------------------------------------------
# build_descriptor — không còn cột vật lý/merge, raw_count luôn = 1
# ---------------------------------------------------------------------------

def _spec(**overrides) -> wio.WizardTableSpec:
    kwargs = dict(
        table_id="A9", name="Bảng thử nghiệm", order=9, value_unit="dBm",
        value_format="dbm",
        rows=[wio.WizardRowSpec(key="f1", freq_set=1e6, limit="≤ 15 mVrms"),
              wio.WizardRowSpec(key="f2", freq_set=2e6, limit="≤ 15 mVrms")],
        pass_rule={"type": "value_vs_parsed_threshold"}, gcn=None,
    )
    kwargs.update(overrides)
    return wio.WizardTableSpec(**kwargs)


def test_build_descriptor_raw_count_always_one():
    d = wio.build_descriptor(_spec())
    assert d.columns == []
    assert d.merge == []
    assert all(r.raw_count == 1 for r in d.rows)


def test_build_descriptor_json_roundtrip(tmp_path):
    spec = _spec()
    d = wio.build_descriptor(spec)
    out = wio.write_descriptor_json(d, tmp_path)
    reloaded = load_table_descriptor(out)
    assert reloaded.table_id == "A9"
    assert reloaded.value_format == "dbm"
    assert reloaded.to_dict() == d.to_dict()


def test_build_descriptor_keeps_gcn_field():
    spec = _spec(gcn={"param_name": "Độ nhạy", "limit_str": "≤ 15 mVrms"})
    d = wio.build_descriptor(spec)
    assert d.gcn == {"param_name": "Độ nhạy", "limit_str": "≤ 15 mVrms"}


# ---------------------------------------------------------------------------
# is_advanced_table / descriptor_to_spec — phân biệt bảng "đơn giản" (form
# Quản lý mẫu báo cáo sửa được) với bảng "nâng cao" (nhiều report_val()/
# dòng — measured_count/value_format_seq/uncertainty_index — chỉ xem JSON).
# ---------------------------------------------------------------------------

def test_is_advanced_table_false_for_plain_descriptor():
    d = wio.build_descriptor(_spec())
    assert wio.is_advanced_table(d) is False


def test_is_advanced_table_true_when_any_row_raw_count_over_1():
    d = wio.build_descriptor(_spec())
    d.rows[0].raw_count = 2
    assert wio.is_advanced_table(d) is True


def test_descriptor_to_spec_round_trips_simple_table():
    original_spec = _spec()
    d = wio.build_descriptor(original_spec)
    spec = wio.descriptor_to_spec(d)
    assert spec.table_id == "A9"
    assert spec.name == "Bảng thử nghiệm"
    assert spec.value_unit == "dBm"
    assert spec.value_format == "dbm"
    assert spec.pass_rule == {"type": "value_vs_parsed_threshold"}
    assert [r.key for r in spec.rows] == ["f1", "f2"]
    assert spec.rows[0].limit == "≤ 15 mVrms"
    # build_descriptor(descriptor_to_spec(d)) phải cho lại ĐÚNG descriptor cũ
    assert wio.build_descriptor(spec).to_dict() == d.to_dict()


# ---------------------------------------------------------------------------
# scan_meta_paragraphs — gợi ý điền sẵn form meta khi tạo mẫu mới
# ---------------------------------------------------------------------------

def test_scan_meta_paragraphs(tmp_path):
    doc = Document()
    doc.add_paragraph("Ký hiệu: CNT-90XL")
    doc.add_paragraph("Không có dấu hai chấm ở đây")
    doc.add_paragraph("Số hiệu: SN12345")
    path = tmp_path / "meta.docx"
    doc.save(str(path))

    pairs = wio.scan_meta_paragraphs(path)
    labels = {p.label: p.value for p in pairs}
    assert labels.get("Ký hiệu") == "CNT-90XL"
    assert labels.get("Số hiệu") == "SN12345"
    assert "Không có dấu hai chấm ở đây" not in labels


# ---------------------------------------------------------------------------
# find_missing_table_ids — cảnh báo nhẹ, chỉ đọc, không sửa file
# ---------------------------------------------------------------------------

def _docx_with_tags(tmp_path, name, tags: list) -> Path:
    doc = Document()
    for t in tags:
        doc.add_paragraph(t)
    path = tmp_path / name
    doc.save(str(path))
    return path


def test_find_missing_table_ids_detects_typo(tmp_path):
    path = _docx_with_tags(tmp_path, "bienban.docx",
                            ["{% if tables.A2.enabled %}", "{{ tables.A2.report_val() }}", "{% endif %}"])
    assert wio.find_missing_table_ids(path, ["A2", "A9"]) == ["A9"]


def test_find_missing_table_ids_all_present(tmp_path):
    path = _docx_with_tags(tmp_path, "bienban.docx",
                            ["{{ tables.A2.report_val() }}", "{{ tables.A9.result }}"])
    assert wio.find_missing_table_ids(path, ["A2", "A9"]) == []


def test_find_missing_table_ids_reads_table_cells_too(tmp_path):
    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    tbl.cell(0, 0).text = "{{ tables.A2.report_val() }}"
    path = tmp_path / "bienban.docx"
    doc.save(str(path))
    assert wio.find_missing_table_ids(path, ["A2"]) == []


# ---------------------------------------------------------------------------
# find_docx_table_grid / value_columns_from_grid — Bước 2 đọc trực tiếp cấu
# trúc cột từ bienban.docx thật để khớp ĐÚNG vị trí report_val()
# ---------------------------------------------------------------------------

def _docx_with_tables(tmp_path, tables: list) -> Path:
    """tables: list[list[list[str]]] — mỗi phần tử là 1 grid (list[list[str]])
    của 1 bảng, ghi thẳng vào .docx theo đúng thứ tự."""
    doc = Document()
    for grid in tables:
        n_rows, n_cols = len(grid), len(grid[0])
        tbl = doc.add_table(rows=n_rows, cols=n_cols)
        for r, row in enumerate(grid):
            for c, text in enumerate(row):
                tbl.cell(r, c).text = text
    path = tmp_path / "bienban.docx"
    doc.save(str(path))
    return path


def test_find_docx_table_grid_locates_correct_table(tmp_path):
    grid_a1 = [["Khoá", "lần 1"], ["10MHz", "{{ tables.A1.report_val() }}"]]
    grid_a2 = [["Khoá", "lần 1", "lần 2"],
               ["1mW", "{{ tables.A2.report_val() }}", "{{ tables.A2.report_val() }}"]]
    path = _docx_with_tables(tmp_path, [grid_a1, grid_a2])

    assert wio.find_docx_table_grid(path, "A2") == grid_a2
    assert wio.find_docx_table_grid(path, "A1") == grid_a1
    assert wio.find_docx_table_grid(path, "A9") is None


def test_find_docx_table_grid_missing_file_returns_none(tmp_path):
    assert wio.find_docx_table_grid(tmp_path / "khong_ton_tai.docx", "A1") is None


def test_value_columns_from_grid_detects_report_val_columns_in_order():
    grid = [
        ["Công suất chuẩn", "lần 1", "lần 2", "Độ KĐBĐ"],
        ["1 mW", "{{ tables.A1.report_val() }}", "{{ tables.A1.report_val() }}", ""],
        ["Trung Bình", "{{ tables.A1.report_val() }}", "{{ tables.A1.report_val() }}",
         "{{ tables.A1.report_val() }}"],
    ]
    assert wio.value_columns_from_grid(grid) == [1, 2, 3]


def test_value_columns_from_grid_empty_grid_returns_empty_list():
    assert wio.value_columns_from_grid([]) == []


# ---------------------------------------------------------------------------
# Ô gộp ngang (gridSpan) — python-docx's row.cells LẶP LẠI cùng 1 ô cho mỗi
# cột nó gridSpan qua (hành vi mặc định, không phải bug của python-docx) —
# scan_docx_tables()/insert_report_val_tags() phải khử trùng lặp này, không
# thì grid ra NHIỀU cột hơn số cột NHÌN THẤY thật trong Word (bug thật đã
# gặp: cột "lần 5" gộp 2 ô lưới bị đọc thành "lần 5" LẶP LẠI 2 lần).
# ---------------------------------------------------------------------------

def _docx_with_merged_header_cell(tmp_path) -> Path:
    """3 cột lưới khai báo trong bảng, nhưng cột 2+3 GỘP NGANG thành 1 ô
    "lần 5" duy nhất (giống bảng thật đã gặp lỗi) -> bảng chỉ có ĐÚNG 2 cột
    NHÌN THẤY thật: "Khoá" và "lần 5"."""
    doc = Document()
    tbl = doc.add_table(rows=2, cols=3)
    tbl.cell(0, 0).text = "Khoá"
    merged_header = tbl.cell(0, 1).merge(tbl.cell(0, 2))
    merged_header.text = "lần 5"
    tbl.cell(1, 0).text = "1 mW"
    merged_data = tbl.cell(1, 1).merge(tbl.cell(1, 2))
    merged_data.text = "{{ tables.A1.report_val() }}"
    path = tmp_path / "bienban.docx"
    doc.save(str(path))
    return path


def test_scan_docx_tables_dedupes_gridspan_merged_cell(tmp_path):
    """Bug thật: ô "lần 5" gộp 2 cột lưới bị đọc thành 2 cột trùng nhau —
    grid phải chỉ có ĐÚNG 2 cột nhìn thấy thật, không phải 3."""
    path = _docx_with_merged_header_cell(tmp_path)
    detected = wio.scan_docx_tables(path)[0]
    assert detected.grid[0] == ["Khoá", "lần 5"]
    assert detected.n_cols == 2


def test_insert_report_val_tags_writes_to_correct_deduped_column(tmp_path):
    """measured_cols đánh số theo cột đã khử trùng lặp (gridSpan) — ghi tag
    phải nhắm ĐÚNG ô đó, không lệch cột do row.cells chưa khử trùng lặp."""
    path = _docx_with_merged_header_cell(tmp_path)
    # Chưa gắn tag ở cột "Khoá" (cột 0) — chỉ gắn cột 1 ("lần 4").
    n = wio.insert_report_val_tags(path, table_index=0, measured_cols=[1], table_id="A2",
                                   header_row_index=0)
    assert n == 1
    detected = wio.scan_docx_tables(path)[0]
    assert detected.grid[1][0] == "1 mW"   # cột Khoá KHÔNG bị đụng tới
    assert detected.grid[1][1] == "{{ tables.A2.report_val() }}"


def _docx_with_summary_row(tmp_path) -> Path:
    """5 cột lưới: Khoá + 4 cột "lần N" ở dòng tiêu đề/dòng dữ liệu thường,
    nhưng dòng tổng hợp "Trung Bình" gộp cả 4 cột "lần N" thành 1 ô rộng
    duy nhất — khác hẳn cấu trúc dòng tiêu đề (giống bảng thật đã gặp lỗi:
    dòng "Trung Bình" trong TEMPLATE_POWER-style, gộp nhiều lần đo lại)."""
    doc = Document()
    tbl = doc.add_table(rows=3, cols=5)
    tbl.cell(0, 0).text = "Khoá"
    for c in range(1, 5):
        tbl.cell(0, c).text = f"lần {c}"
    tbl.cell(1, 0).text = "1 mW"
    for c in range(1, 5):
        tbl.cell(1, c).text = f"raw{c}"
    tbl.cell(2, 0).text = "Trung Bình"
    merged = tbl.cell(2, 1).merge(tbl.cell(2, 4))
    merged.text = "raw_tb"
    path = tmp_path / "bienban.docx"
    doc.save(str(path))
    return path


def test_insert_report_val_tags_writes_one_tag_into_merged_summary_row(tmp_path):
    """KHÔNG còn báo lỗi/crash cho dòng tổng hợp gộp ô khác dòng tiêu đề —
    ghi ĐÚNG 1 tag vào ô rộng đã gộp (khớp đúng số ô THẬT dòng đó có), các
    dòng dữ liệu thường vẫn nhận đủ 4 tag riêng biệt như bình thường."""
    path = _docx_with_summary_row(tmp_path)

    n = wio.insert_report_val_tags(path, table_index=0, measured_cols=[1, 2, 3, 4],
                                   table_id="A1", header_row_index=0)
    assert n == 4 + 1   # 4 tag (dòng "1 mW") + 1 tag (dòng "Trung Bình", đã gộp)

    detected = wio.scan_docx_tables(path)[0]
    assert detected.grid[1][1:] == ["{{ tables.A1.report_val() }}"] * 4
    assert detected.grid[2] == ["Trung Bình", "{{ tables.A1.report_val() }}"]


def test_raw_counts_for_measured_cols_matches_actual_tags_written(tmp_path):
    """raw_counts_for_measured_cols() (dry-run, không ghi gì) phải trả ĐÚNG
    số report_val() mà insert_report_val_tags() thực sự ghi cho từng dòng —
    2 hàm dùng chung logic quy đổi vị trí lưới nên luôn khớp nhau."""
    path = _docx_with_summary_row(tmp_path)
    counts = wio.raw_counts_for_measured_cols(path, table_index=0, measured_cols=[1, 2, 3, 4],
                                              header_row_index=0)
    assert counts == [4, 1]   # dòng "1 mW" = 4, dòng "Trung Bình" (gộp) = 1

    wio.insert_report_val_tags(path, table_index=0, measured_cols=[1, 2, 3, 4],
                               table_id="A1", header_row_index=0)
    detected = wio.scan_docx_tables(path)[0]
    actual_counts = [sum(1 for cell in row if "report_val()" in cell) for row in detected.grid[1:]]
    assert actual_counts == counts


def test_raw_counts_expands_full_span_of_merged_header_cell(tmp_path):
    """Bug thật đã gặp: nếu CHÍNH ô header đã chọn cũng gộp nhiều cột lưới
    (vd "lần 5" gộp 2 cột), phải tính CẢ 2 vị trí lưới nó chiếm — chỉ tính
    vị trí BẮT ĐẦU sẽ làm dòng tổng hợp thiếu report_val() (undercounting)
    khi dòng đó có 1 ô khác BẮT ĐẦU đúng tại vị trí lưới THỨ 2 của ô header
    đã gộp — y hệt bảng A1 Demo_Ky thật (dòng "Trung Bình" từng ra
    raw_count=1 thay vì 2 đúng)."""
    doc = Document()
    # 7 cột lưới: Khoá, lần1, lần2, lần3, lần4, [lần5 gộp 2 cột lưới: 5+6].
    tbl = doc.add_table(rows=2, cols=7)
    tbl.cell(0, 0).text = "Khoá"
    for c in range(1, 5):
        tbl.cell(0, c).text = f"lần {c}"
    merged_header = tbl.cell(0, 5).merge(tbl.cell(0, 6))
    merged_header.text = "lần 5"
    # Dòng "Trung Bình": 1 ô gộp vị trí 1-5 (span=5, BẮT ĐẦU tại vị trí 1 —
    # đã nằm trong measured_positions dù có fix hay không), rồi 1 ô RIÊNG
    # BẮT ĐẦU tại vị trí 6 (span=1) — vị trí 6 CHỈ được tính là "giá trị đo"
    # nếu quy đổi header đúng CẢ span của ô "lần 5" (không chỉ vị trí bắt
    # đầu = 5), đúng điểm bug đã gặp.
    tbl.cell(1, 0).text = "Trung Bình"
    merged_1 = tbl.cell(1, 1).merge(tbl.cell(1, 5))
    merged_1.text = "raw_tb_1"
    tbl.cell(1, 6).text = "raw_tb_2"
    path = tmp_path / "bienban.docx"
    doc.save(str(path))

    # measured_cols theo index header ĐÃ KHỬ TRÙNG LẶP: 0=Khoá,1..4=lần1-4,5=lần5(span2).
    counts = wio.raw_counts_for_measured_cols(path, table_index=0, measured_cols=[1, 2, 3, 4, 5],
                                              header_row_index=0)
    assert counts == [2]   # 1 ô gộp (vị trí 1-5) + 1 ô riêng (vị trí 6) = 2, KHÔNG phải 1


# ---------------------------------------------------------------------------
# extra_skip_rows — bảng có NHIỀU dòng tiêu đề lồng nhau (vd 1 dòng phụ ghi
# lại tên cột giữa chừng bảng, giống bảng A1 Demo_Ky thật: dòng 0 = "lần
# 1-5", dòng phụ ở giữa = "lần 6-10") — dòng phụ đó phải GIỮ NGUYÊN chữ
# tĩnh, KHÔNG bị coi là dữ liệu/KHÔNG bị gắn tag.
# ---------------------------------------------------------------------------

def _docx_with_nested_header_row(tmp_path) -> Path:
    doc = Document()
    tbl = doc.add_table(rows=3, cols=3)
    tbl.cell(0, 0).text = "Khoá"
    tbl.cell(0, 1).text = "lần 1"
    tbl.cell(0, 2).text = "lần 2"
    tbl.cell(1, 0).text = ""
    tbl.cell(1, 1).text = "lần 3"
    tbl.cell(1, 2).text = "lần 4"
    tbl.cell(2, 0).text = "1 mW"
    tbl.cell(2, 1).text = "raw1"
    tbl.cell(2, 2).text = "raw2"
    path = tmp_path / "bienban.docx"
    doc.save(str(path))
    return path


def test_build_rows_from_grid_skips_extra_skip_rows():
    grid = [["Khoá", "lần 1", "lần 2"], ["", "lần 3", "lần 4"], ["1 mW", "raw1", "raw2"]]
    rows = wio.build_rows_from_grid(grid, {0: "display_label"}, header_row_index=0,
                                    extra_skip_rows=frozenset({1}))
    assert [r.key for r in rows] == ["1 mW"]


def test_raw_counts_and_tags_skip_extra_skip_rows(tmp_path):
    path = _docx_with_nested_header_row(tmp_path)
    counts = wio.raw_counts_for_measured_cols(path, table_index=0, measured_cols=[1, 2],
                                              header_row_index=0, extra_skip_rows=frozenset({1}))
    assert counts == [2]   # chỉ đúng 1 dòng dữ liệu thật ("1 mW"), dòng phụ bị bỏ qua

    wio.insert_report_val_tags(path, table_index=0, measured_cols=[1, 2], table_id="A1",
                               header_row_index=0, extra_skip_rows=frozenset({1}))
    detected = wio.scan_docx_tables(path)[0]
    assert detected.grid[1] == ["", "lần 3", "lần 4"]   # dòng phụ GIỮ NGUYÊN, không bị ghi đè
    assert detected.grid[2] == ["1 mW", "{{ tables.A1.report_val() }}", "{{ tables.A1.report_val() }}"]


# ---------------------------------------------------------------------------
# measured_cell_flags() — dùng để CẢNH BÁO trước khi ghi đè (không phải để
# chọn tag qua checkbox từng ô — đã bỏ hướng đó vì bảng xem trước bị rối,
# quay lại chọn theo CỘT ở dòng tiêu đề như cũ, chỉ thêm bước cảnh báo).
# ---------------------------------------------------------------------------

def test_measured_cell_flags_matches_raw_counts_for_measured_cols(tmp_path):
    path = _docx_with_summary_row(tmp_path)
    flags = wio.measured_cell_flags(path, table_index=0, measured_cols=[1, 2, 3, 4], header_row_index=0)
    # dòng 1 ("1 mW"): 4 ô đã khử trùng lặp đều được tick (index 1..4).
    assert flags[1] == {1, 2, 3, 4}
    # dòng 2 ("Trung Bình"): CHỈ 1 ô đã gộp -> đúng 1 index được tick.
    assert flags[2] == {1}
    # dòng tiêu đề (0) vẫn được trả về (khách tự loại ở tầng GUI), không rỗng.
    assert flags[0] == {1, 2, 3, 4}
