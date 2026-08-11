"""
unit_test/test_table_import.py
=================================
Test core/table_import.py — "Quản lý mẫu báo cáo": sao chép 1 mẫu đã có
(copy_template), sửa thông tin chung (update_meta), thay file Word
(replace_docx), và ghi 1 bảng (mới hoặc sửa lại) vào mẫu đã có
(apply_table_to_existing). KHÔNG còn "tạo mẫu rỗng" — mọi mẫu mới đều bắt
đầu từ sao chép 1 mẫu thật.
"""

import json
from pathlib import Path

import pytest
from docx import Document

from core import table_import as timport
from core import table_wizard_io as wio


def _hand_tagged_bienban_docx(path: Path, table_id: str = "A9"):
    """Mô phỏng 1 file Biên Bản mà quản trị viên đã tự gõ tag Jinja trực
    tiếp trong Word — KHÔNG qua cơ chế quét/tự động chèn tag nào cả."""
    doc = Document()
    doc.add_paragraph("Tên phương tiện ĐL-TN: {{ header.name }}")
    doc.add_paragraph("{%% if tables.%s.enabled %%}" % table_id)
    doc.add_paragraph("Bảng %s - Bảng thử" % table_id)
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).text = "Tần số thiết lập"
    tbl.cell(0, 1).text = "Độ nhạy đo được"
    tbl.cell(1, 0).text = "100 kHz"
    tbl.cell(1, 1).text = "{{ tables.%s.report_val() }}" % table_id
    doc.add_paragraph("{% endif %}")
    doc.save(str(path))


def _descriptor(table_id="A9", **overrides):
    kwargs = dict(table_id=table_id, name="Bảng thử", order=1, value_unit="mVrms",
                  value_format="text", rows=[wio.WizardRowSpec(key="100kHz", freq_set=100000.0)],
                  pass_rule={"type": "none"}, gcn=None)
    kwargs.update(overrides)
    return wio.build_descriptor(wio.WizardTableSpec(**kwargs))


def _no_leftover_tags(doc: Document) -> bool:
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    return "{{" not in full_text and "{%" not in full_text


def _seed_template(tpl_dir: Path, template_id: str, bienban_path: Path,
                    descriptor, kind: str = "kiem_dinh", **meta_overrides) -> Path:
    """Dựng 1 thư mục mẫu "từ nguồn" CHỈ bằng thao tác file thuần (mkdir +
    ghi JSON + copy docx) — thay cho apply_new_template() đã bỏ (không còn
    đường "tạo mẫu rỗng" trong app; đây chỉ là fixture cho test cần 1 mẫu
    THẬT làm điểm xuất phát, giống cách app luôn bắt đầu từ mẫu có sẵn)."""
    tables_dir = tpl_dir / "tables"
    tables_dir.mkdir(parents=True, exist_ok=True)
    import shutil
    shutil.copy(str(bienban_path), str(tpl_dir / "bienban.docx"))
    timport.apply_table_to_existing(tables_dir, descriptor)
    meta = {
        "template_id": template_id, "template_name": "Mẫu thử", "kind": kind,
        "dut_models": [], "standard": "", "measurement_range": "",
        "dut_manufacturer_default": "",
    }
    meta.update(meta_overrides)
    (tpl_dir / "meta.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return tpl_dir


# ---------------------------------------------------------------------------
# copy_template — nguồn duy nhất để tạo mẫu MỚI, luôn từ 1 mẫu đã có
# ---------------------------------------------------------------------------

def test_copy_template_duplicates_meta_tables_and_docx(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)

    src_dir = tmp_path / "SRC_TPL"
    bienban_path = tmp_path / "customer_bienban.docx"
    _hand_tagged_bienban_docx(bienban_path)
    _seed_template(src_dir, "SRC_TPL", bienban_path, _descriptor())

    dst_dir = timport.copy_template("SRC_TPL", "DST_TPL", "Mẫu bản sao")

    assert dst_dir == tmp_path / "DST_TPL"
    assert (dst_dir / "tables" / "A9.json").exists()
    assert (dst_dir / "bienban.docx").read_bytes() == bienban_path.read_bytes()
    meta = json.loads((dst_dir / "meta.json").read_text(encoding="utf-8"))
    assert meta["template_id"] == "DST_TPL"
    assert meta["template_name"] == "Mẫu bản sao"


def test_copy_template_rejects_existing_destination(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)

    bienban_path = tmp_path / "customer_bienban.docx"
    _hand_tagged_bienban_docx(bienban_path)
    _seed_template(tmp_path / "SRC_TPL", "SRC_TPL", bienban_path, _descriptor())
    _seed_template(tmp_path / "EXISTS_TPL", "EXISTS_TPL", bienban_path, _descriptor())

    with pytest.raises(ValueError, match="đã tồn tại"):
        timport.copy_template("SRC_TPL", "EXISTS_TPL", "x")


def test_copy_template_rejects_missing_source(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)
    with pytest.raises(ValueError, match="không tồn tại"):
        timport.copy_template("KHONG_TON_TAI", "NEW", "x")


# ---------------------------------------------------------------------------
# delete_template — chuyển vào Thùng rác (send2trash), không xoá vĩnh viễn
# ---------------------------------------------------------------------------

def test_delete_template_sends_whole_dir_to_trash(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)

    bienban_path = tmp_path / "customer_bienban.docx"
    _hand_tagged_bienban_docx(bienban_path)
    tpl_dir = _seed_template(tmp_path / "TPL", "TPL", bienban_path, _descriptor())

    trashed = []
    monkeypatch.setattr("send2trash.send2trash", lambda p: trashed.append(p))

    timport.delete_template("TPL")
    assert trashed == [str(tpl_dir)]


def test_delete_template_rejects_missing_template(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)
    with pytest.raises(ValueError, match="không tồn tại"):
        timport.delete_template("KHONG_TON_TAI")


# ---------------------------------------------------------------------------
# update_meta — sửa thông tin chung, KHÔNG đổi template_id
# ---------------------------------------------------------------------------

def test_update_meta_overwrites_editable_fields_only(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)

    bienban_path = tmp_path / "customer_bienban.docx"
    _hand_tagged_bienban_docx(bienban_path)
    _seed_template(tmp_path / "TPL", "TPL", bienban_path, _descriptor(), standard="Cũ")

    timport.update_meta("TPL", {"template_name": "Tên mới", "standard": "Tiêu chuẩn mới",
                                "dut_models": ["M1", "M2"]})

    meta = json.loads((tmp_path / "TPL" / "meta.json").read_text(encoding="utf-8"))
    assert meta["template_id"] == "TPL"   # không đổi
    assert meta["template_name"] == "Tên mới"
    assert meta["standard"] == "Tiêu chuẩn mới"
    assert meta["dut_models"] == ["M1", "M2"]


# ---------------------------------------------------------------------------
# replace_docx — thay file .docx sống, copy nguyên vẹn
# ---------------------------------------------------------------------------

def test_replace_docx_copies_new_file_byte_identical(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)

    bienban_path = tmp_path / "customer_bienban.docx"
    _hand_tagged_bienban_docx(bienban_path)
    _seed_template(tmp_path / "TPL", "TPL", bienban_path, _descriptor())

    new_docx = tmp_path / "new_bienban.docx"
    _hand_tagged_bienban_docx(new_docx, table_id="A10")

    out = timport.replace_docx("TPL", "bienban", new_docx)
    assert out == tmp_path / "TPL" / "bienban.docx"
    assert out.read_bytes() == new_docx.read_bytes()


# ---------------------------------------------------------------------------
# apply_table_to_existing — ghi JSON (thêm mới HOẶC sửa đè), không đụng docx
# ---------------------------------------------------------------------------

def test_apply_table_to_existing_writes_json_only(tmp_path):
    tables_dir = tmp_path / "tables_out"
    out = timport.apply_table_to_existing(tables_dir=tables_dir, descriptor=_descriptor(table_id="A10"))
    assert out.exists()
    assert out == tables_dir / "A10.json"
    assert list(tmp_path.glob("*.docx")) == []
    assert list(tmp_path.glob("*.bak-*")) == []


def test_apply_table_to_existing_overwrites_when_editing(tmp_path):
    tables_dir = tmp_path / "tables_out"
    timport.apply_table_to_existing(tables_dir, _descriptor(table_id="A10", name="Tên cũ"))
    timport.apply_table_to_existing(tables_dir, _descriptor(table_id="A10", name="Tên mới đã sửa"))

    from core.table_descriptor import load_table_descriptor
    reloaded = load_table_descriptor(tables_dir / "A10.json")
    assert reloaded.name == "Tên mới đã sửa"


# ---------------------------------------------------------------------------
# Round-trip đầy đủ: file .docx TỰ MÔ PHỎNG "khách đã gõ tay" tag -> dựng
# thư mục mẫu bằng thao tác file thuần -> get_template().generate_bienban()
# -> không còn tag Jinja sót + giá trị đúng.
# ---------------------------------------------------------------------------

def test_generic_template_end_to_end_with_hand_typed_tags(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)

    bienban_path = tmp_path / "customer_bienban.docx"
    _hand_tagged_bienban_docx(bienban_path, table_id="A9")
    _seed_template(tmp_path / "TEST_E2E", "TEST_E2E", bienban_path,
                    _descriptor(table_id="A9", value_format="mv"))

    from core.report_templates import get_template
    from core.session import CalibrationSession, SessionTest, TableRow, ReportTable

    tpl = get_template("TEST_E2E")
    tests = tpl.default_tests()
    assert len(tests) == 1 and tests[0].table_id == "A9"

    row = TableRow(key="100kHz", freq_set=100000.0, raw_readings=[12.3], value_unit="mVrms", confirmed=True)
    rt = ReportTable(table_id="A9", name="Bảng thử", rows=[row])
    session = CalibrationSession(template_id="TEST_E2E")
    session.tests = [SessionTest(table_id="A9", name="Bảng thử", enabled=True, result_table=rt)]

    out = tpl.generate_bienban(session, tmp_path / "rendered.docx")
    rendered = Document(str(out))
    assert _no_leftover_tags(rendered)
    full_text = "\n".join(p.text for p in rendered.paragraphs)
    for t in rendered.tables:
        for r in t.rows:
            for c in r.cells:
                full_text += "\n" + c.text
    assert "12,3 mVrms" in full_text


def test_generic_template_end_to_end_disabled_table_hides_section(tmp_path, monkeypatch):
    import core.report_templates.generic as generic_mod
    monkeypatch.setattr(generic_mod, "TEMPLATES_DIR", tmp_path)

    bienban_path = tmp_path / "customer_bienban.docx"
    _hand_tagged_bienban_docx(bienban_path, table_id="A9")
    _seed_template(tmp_path / "TEST_E2E_OFF", "TEST_E2E_OFF", bienban_path, _descriptor(table_id="A9"))

    from core.report_templates import get_template
    from core.session import CalibrationSession, SessionTest

    tpl = get_template("TEST_E2E_OFF")
    session = CalibrationSession(template_id="TEST_E2E_OFF")
    session.tests = [SessionTest(table_id="A9", name="Bảng thử", enabled=False, result_table=None)]

    out = tpl.generate_bienban(session, tmp_path / "rendered_off.docx")
    rendered = Document(str(out))
    assert _no_leftover_tags(rendered)
    has_heading = any("Bảng A9 - Bảng thử" in p.text for p in rendered.paragraphs)
    assert has_heading is False
