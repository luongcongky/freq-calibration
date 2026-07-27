"""
core/report_generator.py
========================
Sinh tài liệu Word (.docx) theo mẫu QTKĐ 2.461 : 2018 từ CalibrationSession.

Hai loại tài liệu:
  • generate_bienban()  → Biên Bản Kiểm Định (Phụ lục A)
  • generate_gcnkd()    → Giấy Chứng Nhận Kiểm Định (Phụ lục B)

Định dạng bám sát PDF mẫu của khách hàng:
  - Font: Times New Roman
  - Cỡ chữ: 12pt (nội dung), 13pt (tiêu đề)
  - Lề: trên 2cm, dưới 2cm, trái 3cm, phải 2cm (chuẩn văn bản nhà nước VN)
  - Bảng: viền mỏng, header in đậm căn giữa
"""

from __future__ import annotations

import math
import logging
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.session import CalibrationSession, ReportTable, TableRow

log = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Hằng số định dạng
# ---------------------------------------------------------------------------
FONT_NAME = "Times New Roman"
SIZE_NORMAL = Pt(12)
SIZE_TITLE  = Pt(13)
SIZE_SMALL  = Pt(11)

_SUP = str.maketrans("0123456789-", "⁰¹²³⁴⁵⁶⁷⁸⁹⁻")


# ---------------------------------------------------------------------------
# Hàm tiện ích định dạng số
# ---------------------------------------------------------------------------

def _sci(value: float, dec: int = 1) -> str:
    """2.4e-7 → '2,4×10⁻⁷' dùng Unicode superscript."""
    if value == 0:
        return "0"
    sign = "-" if value < 0 else ""
    v = abs(value)
    exp = int(math.floor(math.log10(v)))
    m = v / (10 ** exp)
    m_r = round(m, dec)
    if m_r >= 10:
        m_r /= 10; exp += 1
    m_str = f"{m_r:.{dec}f}".replace(".", ",")
    return f"{sign}{m_str}×10{str(exp).translate(_SUP)}"


def _fmt_freq(hz: float) -> str:
    """10000000 → '10 MHz', 1500 → '1,5 kHz'"""
    if hz == 0:
        return "0 Hz"
    if hz < 1e3:
        v = hz
        u = "Hz"
    elif hz < 1e6:
        v = hz / 1e3
        u = "kHz"
    elif hz < 1e9:
        v = hz / 1e6
        u = "MHz"
    else:
        v = hz / 1e9
        u = "GHz"
    s = f"{v:g}".replace(".", ",")
    return f"{s} {u}"


def _fmt_hz_measured(hz: float) -> str:
    """Định dạng giá trị tần số đo được: 9999999.98765 → '9.999.999,98765 Hz'"""
    int_part = int(abs(hz))
    frac = abs(hz) - int_part
    # Nhóm nghìn bằng dấu chấm
    int_str = f"{int_part:,}".replace(",", ".")
    if frac > 0:
        frac_str = f"{frac:.7f}"[1:].rstrip("0")  # ".9876500" → ",98765"
        frac_str = frac_str.replace(".", ",")
        return f"{int_str}{frac_str} Hz"
    return f"{int_str} Hz"


def _fmt_period(s: float) -> str:
    """Định dạng chu kỳ đo được."""
    if s == 0:
        return "0 s"
    if s >= 1e-1:
        return f"{s*1e3:.6g} ms".replace(".", ",")
    if s >= 1e-4:
        return f"{s*1e6:.6g} µs".replace(".", ",")
    if s >= 1e-7:
        return f"{s*1e9:.6g} ns".replace(".", ",")
    return f"{s:.4e} s".replace(".", ",").replace("e-0", "×10⁻").replace("e-", "×10⁻")


def _fmt_mv(mv: float) -> str:
    """14.9 → '14,9 mVrms'"""
    return f"{mv:.4g} mVrms".replace(".", ",")


def _fmt_dbm(dbm: float) -> str:
    """-35.1 → '-35,1 dBm'"""
    return f"{dbm:.2f} dBm".replace(".", ",")


def _pass_mark(passed: Optional[bool]) -> str:
    if passed is True:
        return "Đạt"
    if passed is False:
        return "Không đạt"
    return ""


# ---------------------------------------------------------------------------
# Tiện ích python-docx
# ---------------------------------------------------------------------------

def _set_font(run, size=SIZE_NORMAL, bold=False, italic=False, underline=False,
              color: Optional[RGBColor] = None):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = color


def _para(doc: Document, text: str = "", align=WD_ALIGN_PARAGRAPH.LEFT,
          size=SIZE_NORMAL, bold=False, italic=False, underline=False,
          space_before=0, space_after=0) -> object:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic, underline=underline)
    return p


def _cell_para(cell, text: str = "", align=WD_ALIGN_PARAGRAPH.CENTER,
               bold=False, size=SIZE_NORMAL, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.clear()
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic)
    return p


def _set_col_widths(table, widths_cm: list[float]):
    """Đặt độ rộng cột (cm). Phải gọi sau khi tạo bảng."""
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])
    # Lock layout = fixed để Word không tự điều chỉnh
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    existing = tblPr.find(qn("w:tblLayout"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblLayout)


def _merge_col(table, col: int, row_start: int, row_end: int):
    """Gộp các ô trong cùng một cột từ row_start đến row_end (inclusive)."""
    start_cell = table.cell(row_start, col)
    end_cell   = table.cell(row_end,   col)
    start_cell.merge(end_cell)


def _page_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(__import__("docx.enum.text", fromlist=["WD_BREAK"]).WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Header tài liệu theo chuẩn văn bản nhà nước Việt Nam
# ---------------------------------------------------------------------------

def _add_gov_header(doc: Document, session: CalibrationSession):
    """
    Thêm phần đầu trang kiểu văn bản nhà nước VN (2 cột):
      Trái: Cục Tiêu chuẩn / Trung tâm ...
      Phải: Cộng hoà XHCN Việt Nam / Độc lập...
    """
    # Dùng bảng 2 cột, không có viền
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    # Xóa viền bảng
    for row in tbl.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)

    left  = tbl.cell(0, 0)
    right = tbl.cell(0, 1)
    left.width  = Cm(8)
    right.width = Cm(8)

    # --- Cột trái ---
    lp = left.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(lp, "CỤC TIÊU CHUẨN\n", size=SIZE_NORMAL)
    _add_run(lp, "ĐO LƯỜNG-CHẤT LƯỢNG\n", size=SIZE_NORMAL)
    _add_run(lp, "TRUNG TÂM TIÊU CHUẨN\nĐO LƯỜNG-CHẤT LƯỢNG 2", size=SIZE_NORMAL, bold=True)
    # Gạch dưới tên đơn vị
    lp2 = left.add_paragraph("_" * 20)
    lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Cột phải ---
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(rp, "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n", size=SIZE_NORMAL, bold=True)
    _add_run(rp, "Độc lập - Tự do - Hạnh phúc", size=SIZE_NORMAL, bold=True, underline=True)
    rp2 = right.add_paragraph("")
    rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp2.paragraph_format.space_before = Pt(2)


def _add_run(para, text: str, size=SIZE_NORMAL, bold=False, italic=False, underline=False):
    run = para.add_run(text)
    _set_font(run, size=size, bold=bold, italic=italic, underline=underline)
    return run


def _clear_cell_borders(cell):
    """Xóa viền của một ô trong bảng (dùng cho bảng layout không viền)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tcBorders.append(border)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# Thiết lập trang
# ---------------------------------------------------------------------------

def _setup_page(doc: Document):
    """Lề chuẩn văn bản nhà nước VN: T=2, D=2, T=3, P=2 (cm)."""
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)


# ---------------------------------------------------------------------------
# Bảng A1 — Sai số bộ dao động thạch anh
# ---------------------------------------------------------------------------

def _add_table_a1(doc: Document, rt: ReportTable):
    _para(doc, "Bảng A1 - Xác định sai số tần số bộ dao động thạch anh",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_SMALL,
          space_before=4, space_after=2)

    row0 = rt.rows[0] if rt.rows else None
    raw_count = len(row0.raw_readings) if row0 else 10
    n_data = max(raw_count, 1)
    n_rows = 1 + n_data   # header + data rows

    tbl = doc.add_table(rows=n_rows, cols=5)
    tbl.style = "Table Grid"
    widths = [3.2, 4.0, 3.8, 2.5, 2.5]
    _set_col_widths(tbl, widths)

    # Header
    headers = [
        "Tần số\nthiết lập",
        "Tần số đo được\ntrên CNT-90XL\n(fCi)",
        "Tần số đo\nđược trên\nCNT-90XL\n(fC)",
        "Sai số tần số\n(δf)",
        "Sai số\ncho phép\n(δfcp)",
    ]
    for j, h in enumerate(headers):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    # Dữ liệu
    if row0:
        f_set_str = _fmt_freq(row0.freq_set or 10e6)
        f_avg_str = _fmt_hz_measured(row0.value_measured) if row0.value_measured else ""
        err_str   = f"± {_sci(row0.error)}" if row0.error is not None else ""
        raws      = row0.raw_readings

        for i in range(n_data):
            raw_str = _fmt_hz_measured(raws[i]) if i < len(raws) else ""
            _cell_para(tbl.cell(i + 1, 1), raw_str, size=SIZE_SMALL)

        # Gộp cột 0 (tần số thiết lập), 2 (fC), 3 (δf), 4 (giới hạn)
        for col_idx in [0, 2, 3, 4]:
            _merge_col(tbl, col_idx, 1, n_data)

        _cell_para(tbl.cell(1, 0), f_set_str, size=SIZE_SMALL)
        _cell_para(tbl.cell(1, 2), f_avg_str, size=SIZE_SMALL)
        _cell_para(tbl.cell(1, 3), err_str, size=SIZE_SMALL)
        _cell_para(tbl.cell(1, 4), "± 2,4×10⁻⁷", size=SIZE_SMALL)
    else:
        # Bảng trống (chưa chạy)
        _merge_col(tbl, 0, 1, n_data)
        _cell_para(tbl.cell(1, 0), "10 MHz", size=SIZE_SMALL)
        _merge_col(tbl, 2, 1, n_data)
        _merge_col(tbl, 3, 1, n_data)
        _merge_col(tbl, 4, 1, n_data)
        _cell_para(tbl.cell(1, 4), "± 2,4×10⁻⁷", size=SIZE_SMALL)


# ---------------------------------------------------------------------------
# Bảng A2 / A3 — Độ nhạy kênh A / B (mVrms)
# ---------------------------------------------------------------------------

def _add_table_sensitivity_mv(doc: Document, rt: ReportTable, label: str):
    _para(doc, f"Bảng {rt.table_id} - {rt.name}",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_SMALL,
          space_before=4, space_after=2)

    n_rows = 1 + len(rt.rows)
    tbl = doc.add_table(rows=n_rows, cols=3)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [4.5, 5.0, 6.5])

    for j, h in enumerate(["Tần số thiết lập", "Độ nhạy đo được", "Độ nhạy cho phép"]):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    # Nhóm các hàng có cùng "Độ nhạy cho phép" để gộp ô
    _fill_sensitivity_rows(tbl, rt.rows, unit="mVrms")


def _fill_sensitivity_rows(tbl, rows: list, unit: str):
    """Điền dữ liệu và gộp ô cột 'Độ nhạy cho phép' theo nhóm limit."""
    if not rows:
        return
    # Gộp ô giới hạn theo nhóm liên tiếp cùng limit
    groups: list[tuple[int, int, str]] = []  # (start_row, end_row, limit)
    cur_limit = rows[0].limit
    cur_start = 1
    for i, row in enumerate(rows):
        if row.limit != cur_limit:
            groups.append((cur_start, i, cur_limit))
            cur_start = i + 1
            cur_limit = row.limit
    groups.append((cur_start, len(rows), cur_limit))

    for i, row in enumerate(rows, start=1):
        if row.freq_set:
            _cell_para(tbl.cell(i, 0), _fmt_freq(row.freq_set), size=SIZE_SMALL)
        if row.value_measured is not None:
            if unit == "mVrms":
                _cell_para(tbl.cell(i, 1), _fmt_mv(row.value_measured), size=SIZE_SMALL)
            else:
                _cell_para(tbl.cell(i, 1), _fmt_dbm(row.value_measured), size=SIZE_SMALL)

    for start, end, limit in groups:
        if end > start:
            _merge_col(tbl, 2, start, end)
        _cell_para(tbl.cell(start, 2), limit, size=SIZE_SMALL)


# ---------------------------------------------------------------------------
# Bảng A4 — Độ nhạy kênh C (dBm)
# ---------------------------------------------------------------------------

def _add_table_sensitivity_dbm(doc: Document, rt: ReportTable):
    _para(doc, f"Bảng A4 - {rt.name}",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_SMALL,
          space_before=4, space_after=2)

    n_rows = 1 + len(rt.rows)
    tbl = doc.add_table(rows=n_rows, cols=3)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [4.5, 5.0, 6.5])

    for j, h in enumerate(["Tần số thiết lập", "Độ nhạy đo được", "Độ nhạy cho phép"]):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    _fill_sensitivity_rows(tbl, rt.rows, unit="dBm")


# ---------------------------------------------------------------------------
# Bảng A5 / A6 / A7 — Sai số đo tần số
# ---------------------------------------------------------------------------

def _add_table_freq_error(doc: Document, rt: ReportTable, section_label: str):
    _para(doc, f"Bảng {rt.table_id} - {rt.name}",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_SMALL,
          space_before=4, space_after=2)

    n_rows = 1 + len(rt.rows)
    tbl = doc.add_table(rows=n_rows, cols=4)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [3.5, 5.5, 3.5, 3.5])

    for j, h in enumerate([
        "Tần số\nthiết lập",
        "Tần số đo được\ntrên kênh " + section_label + " (fđo)",
        "Sai số đo tần\nsố (δf)",
        "Sai số\ncho phép",
    ]):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    # Cột giới hạn gộp tất cả data rows
    if rt.rows:
        _merge_col(tbl, 3, 1, len(rt.rows))
        _cell_para(tbl.cell(1, 3), "± 2,4×10⁻⁷", size=SIZE_SMALL)

    for i, row in enumerate(rt.rows, start=1):
        if row.freq_set:
            _cell_para(tbl.cell(i, 0), _fmt_freq(row.freq_set), size=SIZE_SMALL)
        if row.value_measured:
            _cell_para(tbl.cell(i, 1), _fmt_hz_measured(row.value_measured), size=SIZE_SMALL)
        if row.error is not None:
            _cell_para(tbl.cell(i, 2), _sci(row.error), size=SIZE_SMALL)


# ---------------------------------------------------------------------------
# Bảng A8 — Sai số đo chu kỳ
# ---------------------------------------------------------------------------

def _add_table_period_error(doc: Document, rt: ReportTable):
    _para(doc, "Bảng A8 - Xác định sai số đo chu kỳ",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_SMALL,
          space_before=4, space_after=2)

    n_rows = 1 + len(rt.rows)
    tbl = doc.add_table(rows=n_rows, cols=4)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [4.5, 4.5, 3.0, 4.0])

    for j, h in enumerate([
        "Tần số (chu kỳ) thiết lập\ntrên Γ3-110, SMF-100A",
        "Chu kỳ\nđo được (Tđo)",
        "Sai số đo\n(δT)",
        "Sai số cho phép\n(δTcp)",
    ]):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    if rt.rows:
        _merge_col(tbl, 3, 1, len(rt.rows))
        _cell_para(tbl.cell(1, 3), "± 2,4×10⁻⁷", size=SIZE_SMALL)

    for i, row in enumerate(rt.rows, start=1):
        _cell_para(tbl.cell(i, 0), row.key, size=SIZE_SMALL)  # "5 Hz (200 ms)"
        if row.value_measured:
            _cell_para(tbl.cell(i, 1), _fmt_period(row.value_measured), size=SIZE_SMALL)
        if row.error is not None:
            _cell_para(tbl.cell(i, 2), _sci(row.error), size=SIZE_SMALL)


# ---------------------------------------------------------------------------
# Tìm ReportTable theo table_id trong session
# ---------------------------------------------------------------------------

def _find_table(session: CalibrationSession, table_id: str) -> Optional[ReportTable]:
    for test in session.tests:
        if test.table_id == table_id and test.result_table:
            return test.result_table
    return None


def _confirmed(rt: Optional[ReportTable]) -> Optional[ReportTable]:
    """Chỉ giữ lại các dòng đã được người dùng rà soát & xác nhận đưa vào báo cáo."""
    if rt is None:
        return None
    return ReportTable(table_id=rt.table_id, name=rt.name,
                       rows=rt.confirmed_rows(), passed=rt.passed)


def _empty_table(table_id: str, name: str) -> ReportTable:
    """Bảng trống (dùng khi bài test chưa chạy)."""
    from core.session import ReportTable as RT
    return RT(table_id=table_id, name=name, rows=[], passed=None)


# ---------------------------------------------------------------------------
# BIÊN BẢN KIỂM ĐỊNH (Phụ lục A)
# ---------------------------------------------------------------------------

def generate_bienban(session: CalibrationSession, output_path: str | Path) -> Path:
    """
    Sinh Biên Bản Kiểm Định theo mẫu QTKĐ 2.461 : 2018 (Phụ lục A).
    Trả đường dẫn file đã lưu.
    """
    output_path = Path(output_path)
    doc = Document()
    _setup_page(doc)

    # Xóa đoạn trắng mặc định
    for p in doc.paragraphs:
        p.clear()

    meta = session.meta
    dut  = meta.dut

    # ---- Header nhà nước ----
    _add_gov_header(doc, session)
    _para(doc, "", space_after=4)

    # ---- Tiêu đề ----
    _para(doc, "BIÊN BẢN KIỂM ĐỊNH",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_TITLE,
          space_before=6, space_after=4)

    # ---- Thông tin chung ----
    def _info_line(label: str, value: str, tab_pos: float = 0):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(label)
        _set_font(r1, bold=False)
        r2 = p.add_run(value)
        _set_font(r2)
        return p

    def _info_line_2col(label1, val1, label2, val2):
        """Hai trường trên cùng một dòng (dùng tab)."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        _add_run(p, label1); _add_run(p, val1)
        _add_run(p, "          ")
        _add_run(p, label2); _add_run(p, val2)

    _info_line("Tên phương tiện ĐL-TN: ", "Máy đếm tần số")
    _info_line_2col("Ký hiệu: ", dut.model or "CNT-90XL", "Số hiệu: ", dut.serial)
    _info_line("Nước (hãng) sản xuất: ", dut.manufacturer or "Pendulum")
    _info_line("Đặc tính đo lường: ", f"Dải tần số đo từ {dut.measurement_range or '0,002 Hz đến 27 GHz'}")
    _info_line("Phương pháp kiểm định: ", "QTKĐ 2.461 : 2018")
    _info_line("Phương tiện kiểm định: ", meta.inspection_equipment)

    # Môi trường — 2 trường trên 1 dòng
    _info_line_2col("Điều kiện môi trường:  nhiệt độ ", meta.temperature,
                    "      độ ẩm ", meta.humidity)

    # Ngày kiểm định
    if meta.date:
        d = meta.date
        _info_line("Đã tiến hành kiểm định ngày ",
                   f"{d.day:02d} tháng {d.month:02d} năm {d.year}")
    else:
        _info_line("Đã tiến hành kiểm định ngày ", "        tháng         năm 20")

    _para(doc, "KẾT QUẢ KIỂM ĐỊNH",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_NORMAL,
          space_before=6, space_after=4)

    # ---- Mục 1, 2 ----
    _para(doc, "1 Kiểm tra bên ngoài", bold=True, space_before=2, space_after=1)
    _para(doc, "2 Kiểm tra kỹ thuật",  bold=True, space_before=2, space_after=1)
    _para(doc, "3 Kiểm tra đo lường",  bold=True, space_before=2, space_after=2)

    # ---- Các bảng kiểm tra đo lường ----
    _add_bienban_tables(doc, session)

    # ---- Kết luận ----
    _para(doc, "", space_after=4)
    kl = doc.add_paragraph()
    kl.paragraph_format.space_before = Pt(4)
    _add_run(kl, "4 Kết luận: ", bold=True)
    _add_run(kl, meta.conclusion or "")

    # ---- Chữ ký ----
    _para(doc, "", space_after=8)
    tbl_sign = doc.add_table(rows=1, cols=2)
    tbl_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl_sign.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)

    lc = tbl_sign.cell(0, 0)
    rc = tbl_sign.cell(0, 1)
    for cell in (lc, rc):
        cell.width = Cm(8)

    lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(lp, "Người soát lại", bold=True)
    lp2 = lc.add_paragraph(meta.reviewer or ""); lp2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(rp, "Kiểm định viên", bold=True)
    rp2 = rc.add_paragraph(meta.operator or ""); rp2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
    log.info("Đã xuất Biên Bản: %s", output_path)
    return output_path


def _add_bienban_tables(doc: Document, session: CalibrationSession):
    """Thêm các bảng vào Biên Bản — chỉ bảng của bài test được bật (enabled)."""
    enabled = {t.table_id for t in session.tests if t.enabled}

    if "A1" in enabled:
        _para(doc, "3.1 Xác định sai số tần số bộ dao động thạch anh",
              bold=True, space_before=4, space_after=2)
        _add_table_a1(doc, _confirmed(_find_table(session, "A1")) or _empty_table("A1", ""))

    if "A2" in enabled:
        _para(doc, "3.2 Xác định độ nhạy đầu vào kênh A",
              bold=True, space_before=6, space_after=2)
        _add_table_sensitivity_mv(
            doc, _confirmed(_find_table(session, "A2")) or _empty_table("A2", "Xác định độ nhạy đầu vào kênh A"),
            "kênh A")

    if "A3" in enabled:
        _para(doc, "3.3 Xác định độ nhạy đầu vào kênh B",
              bold=True, space_before=6, space_after=2)
        _add_table_sensitivity_mv(
            doc, _confirmed(_find_table(session, "A3")) or _empty_table("A3", "Xác định độ nhạy đầu vào kênh B"),
            "kênh B")

    if "A4" in enabled:
        _para(doc, "3.4 Xác định độ nhạy đầu vào kênh C",
              bold=True, space_before=6, space_after=2)
        _add_table_sensitivity_dbm(
            doc, _confirmed(_find_table(session, "A4")) or _empty_table("A4", "Xác định độ nhạy đầu vào kênh C"))

    if "A5" in enabled:
        _para(doc, "3.5 Xác định sai số đo tần số kênh A",
              bold=True, space_before=6, space_after=2)
        _add_table_freq_error(
            doc, _confirmed(_find_table(session, "A5")) or _empty_table("A5", "Xác định sai số đo tần số kênh A"),
            "A")

    if "A6" in enabled:
        _para(doc, "3.6 Xác định sai số đo tần số kênh B",
              bold=True, space_before=6, space_after=2)
        _add_table_freq_error(
            doc, _confirmed(_find_table(session, "A6")) or _empty_table("A6", "Xác định sai số đo tần số kênh B"),
            "B")

    if "A7" in enabled:
        _para(doc, "3.7 Xác định sai số đo tần số kênh C",
              bold=True, space_before=6, space_after=2)
        _add_table_freq_error(
            doc, _confirmed(_find_table(session, "A7")) or _empty_table("A7", "Xác định sai số đo tần số kênh C"),
            "C")

    if "A8" in enabled:
        _para(doc, "3.8 Xác định sai số đo chu kỳ",
              bold=True, space_before=6, space_after=2)
        _add_table_period_error(
            doc, _confirmed(_find_table(session, "A8")) or _empty_table("A8", "Xác định sai số đo chu kỳ"))


# ---------------------------------------------------------------------------
# GIẤY CHỨNG NHẬN KIỂM ĐỊNH (Phụ lục B)
# ---------------------------------------------------------------------------

def generate_gcnkd(session: CalibrationSession, output_path: str | Path) -> Path:
    """
    Sinh Giấy Chứng Nhận Kiểm Định theo mẫu QTKĐ 2.461 : 2018 (Phụ lục B).
    """
    output_path = Path(output_path)
    doc = Document()
    _setup_page(doc)
    for p in doc.paragraphs:
        p.clear()

    meta = session.meta
    dut  = meta.dut

    # ---- Header ----
    _add_gov_header(doc, session)

    # Ngày tháng phải
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if meta.date:
        d = meta.date
        loc = meta.location or "Thành phố Hồ Chí Minh"
        _add_run(p_date, f"{loc}, ngày {d.day:02d} tháng {d.month:02d} năm {d.year}", italic=True)
    _para(doc, "", space_after=2)

    # ---- Tiêu đề ----
    _para(doc, "GIẤY CHỨNG NHẬN KIỂM ĐỊNH",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_TITLE,
          space_before=4, space_after=2)
    _para(doc, f"Số: {meta.cert_number}",
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    # ---- Thông tin thiết bị ----
    def _line(lbl, val):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        _add_run(p, lbl); _add_run(p, val)

    def _line2(lbl1, v1, lbl2, v2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        _add_run(p, lbl1); _add_run(p, v1)
        _add_run(p, "     ")
        _add_run(p, lbl2); _add_run(p, v2)

    _line("Tên phương tiện ĐL-TN: ", "Máy đếm tần số")
    _line2("Ký hiệu: ", dut.model or "CNT-90XL", "Số hiệu: ", dut.serial)
    _line2("Nước (hãng) sản xuất: ", dut.manufacturer or "Pendulum",
           "Năm sản xuất: ", "")
    _line("Đơn vị sử dụng: ", dut.owner)
    _line("Đặc tính đo lường: ",
          f"Dải tần số đo từ {dut.measurement_range or '0,002 Hz đến 27 GHz'}.")

    _para(doc, "KẾT QUẢ KIỂM ĐỊNH",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_before=6, space_after=4)

    # ---- Bảng kết quả tổng hợp ----
    _add_gcnkd_summary_table(doc, session)

    # ---- Phương pháp + Kết luận ----
    _para(doc, "", space_after=4)
    _line("Phương pháp kiểm định: ", "QTKĐ 2.461 : 2018")
    p_kl = doc.add_paragraph()
    _add_run(p_kl, "Kết luận: ", bold=True)
    _add_run(p_kl, meta.conclusion or "Đạt yêu cầu kỹ thuật đo lường")

    if meta.valid_until:
        _line("Hiệu lực đến ", meta.valid_until_str() + ".*")
    else:
        _line("Hiệu lực đến ", "          /          .*")

    # ---- Chữ ký 3 cột ----
    _para(doc, "", space_after=8)
    tbl_s = doc.add_table(rows=1, cols=3)
    tbl_s.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl_s.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)
            cell.width = Cm(5.3)

    titles = ["Người kiểm soát\n(Chữ ký, họ tên)",
              "Kiểm định viên\n(Chữ ký, họ tên)",
              "THỦ TRƯỞNG ĐƠN VỊ\n(Ký tên, đóng dấu)"]
    for j, title in enumerate(titles):
        cell = tbl_s.cell(0, j)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(cp, title, bold=True, size=SIZE_SMALL)

    # Note cuối
    _para(doc, "", space_after=4)
    p_note = doc.add_paragraph()
    _add_run(p_note, "(*)", italic=True, size=SIZE_SMALL)
    _add_run(p_note, "Với điều kiện tôn trọng các nguyên tắc sử dụng và bảo quản",
             italic=True, size=SIZE_SMALL)

    doc.save(str(output_path))
    log.info("Đã xuất Giấy chứng nhận: %s", output_path)
    return output_path


def _add_gcnkd_summary_table(doc: Document, session: CalibrationSession):
    """Bảng tổng hợp KẾT QUẢ KIỂM ĐỊNH trong Giấy chứng nhận."""
    _all_rows = [
        ("A1", "1.Xác định sai số bộ dao động thạch anh", "± 2,4×10⁻⁷"),
        ("A2", "2.Xác định độ nhạy đầu vào kênh A",       "Theo QTKĐ"),
        ("A3", "3.Xác định độ nhạy đầu vào kênh B",       "Theo QTKĐ"),
        ("A4", "4.Xác định độ nhạy đầu vào kênh C",       "Theo QTKĐ"),
        ("A5", "5.Xác định sai số đo tần số kênh A",      "± 2,4×10⁻⁷"),
        ("A6", "6.Xác định sai số đo tần số kênh B",      "± 2,4×10⁻⁷"),
        ("A7", "7.Xác định sai số đo tần số kênh C",      "± 2,4×10⁻⁷"),
        ("A8", "8.Xác định sai số đo chu kỳ",             "± 2,4×10⁻⁷"),
    ]
    enabled = {t.table_id for t in session.tests if t.enabled}
    rows_def = [(tid, name, lim) for tid, name, lim in _all_rows if tid in enabled]
    if not rows_def:
        return

    tbl = doc.add_table(rows=1 + len(rows_def), cols=3)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [7.5, 3.5, 5.0])

    for j, h in enumerate(["THAM SỐ KIỂM ĐỊNH", "KẾT QUẢ ĐO", "GIÁ TRỊ CHO PHÉP"]):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    for i, (tid, param_name, limit) in enumerate(rows_def, start=1):
        _cell_para(tbl.cell(i, 0), param_name, align=WD_ALIGN_PARAGRAPH.LEFT, size=SIZE_SMALL)
        rt = _find_table(session, tid)
        result_str = _pass_mark(rt.confirmed_passed) if rt else ""
        _cell_para(tbl.cell(i, 1), result_str, size=SIZE_SMALL)
        _cell_para(tbl.cell(i, 2), limit, size=SIZE_SMALL)
