"""
core/table_engine.py
=====================
Engine CHUNG đọc TableDescriptor (core/table_descriptor.py), dùng bởi mọi
template (core/report_templates/generic.py). 3 việc chính:

  1. map_table()          — StepResult -> ReportTable
  2. build_table_context() — list[TableRow] -> dict Jinja
  3. render_with_table_contexts()/render_gcnkd_summary() — render .docx đầy
     đủ (dựng context + gộp ô hậu-render dựa trên marker ẩn).

Không phụ thuộc Qt → test độc lập.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from docxtpl import DocxTemplate
from docx.shared import Pt

from core.session import CalibrationSession, ReportTable, TableRow
from core.table_descriptor import TableDescriptor
from core.report_generator import (
    _fmt_freq, _fmt_hz_measured, _fmt_period, _fmt_mv, _fmt_dbm, _sci,
    _merge_col, _cell_para, _pass_mark, SIZE_SMALL,
)
from core.report_generator_nrp2 import _fmt_w, _fmt_correction, _power_set_from_key

# ---------------------------------------------------------------------------
# Formula layer — công thức tính lại error/passed từ 1 giá trị đo. Giữ
# NGUYÊN chữ ký các hàm recompute_* công khai (gui/report_preview*.py đang
# import trực tiếp để tính lại khi kiểm định viên sửa tay 1 giá trị) — chỉ
# đổi CÁCH triển khai bên trong để dùng chung với engine mới.
# ---------------------------------------------------------------------------

_FREQ_LIMIT = 2.4e-7   # Giới hạn sai số mặc định theo QTKĐ 2.461 : 2018 — dùng
                       # cho các hàm recompute_* legacy (không có tham số ngoài).


def _relative_error_generic(measured: Optional[float], reference: float,
                             fixed_limit: float) -> tuple:
    """delta = |đo - chuẩn| / chuẩn, so với fixed_limit -> (delta, passed)."""
    delta = abs(measured - reference) / reference if measured is not None else None
    passed = delta is not None and abs(delta) <= fixed_limit
    return delta, (passed if delta is not None else None)


def _relative_error(measured: Optional[float], reference: float) -> tuple:
    return _relative_error_generic(measured, reference, _FREQ_LIMIT)


def recompute_a1(freq_set: float, raw_readings: list) -> tuple:
    """(f_avg, delta_f, passed) từ danh sách lần đo — dùng cho bảng A1."""
    f_avg = sum(raw_readings) / len(raw_readings) if raw_readings else None
    delta_f, passed = _relative_error(f_avg, freq_set)
    return f_avg, delta_f, passed


def recompute_freq_error(freq_set: float, f_meas: Optional[float]) -> tuple:
    """(delta_f, passed) — dùng cho A5/A6/A7."""
    return _relative_error(f_meas, freq_set)


def recompute_period_error(period_set: float, t_meas: Optional[float]) -> tuple:
    """(delta_t, passed) — dùng cho A8."""
    return _relative_error(t_meas, period_set)


def _parse_le_limit(limit_str: str) -> Optional[float]:
    """Phân tích ngưỡng dạng '≤ 15 mVrms' / '≤ -27 dBm' -> số ngưỡng (giữ dấu)."""
    s = limit_str.replace("≤", "").strip()
    for unit in (" mVrms", " dBm"):
        if s.endswith(unit):
            s = s[: -len(unit)]
            break
    try:
        return float(s.replace(",", "."))
    except ValueError:
        return None


def recompute_sensitivity(value: Optional[float], limit_str: str) -> Optional[bool]:
    """passed từ 1 giá trị độ nhạy (mVrms hoặc dBm) so với ngưỡng '≤ ...' —
    dùng cho A2/A3/A4."""
    if value is None:
        return None
    limit = _parse_le_limit(limit_str)
    return limit is not None and value <= limit


def recompute_avg(raw_readings: list, power_set) -> tuple:
    """(giá trị trung bình, số hiệu chỉnh = chuẩn − trung bình) — NRP2."""
    avg = sum(raw_readings) / len(raw_readings) if raw_readings else None
    correction = (power_set - avg) if avg is not None else None
    return avg, correction


# ---------------------------------------------------------------------------
# Generic mapper — StepResult -> ReportTable
# ---------------------------------------------------------------------------

def _table_values(step_results) -> list:
    """Lấy TUẦN TỰ (đúng thứ tự thực thi) mọi giá trị report_val — step_results
    của 1 bài test luôn chỉ ứng với đúng 1 bảng, không cần lọc theo tên bảng."""
    return [r.value for r in step_results
            if r.action == "report_val" and r.ok and r.value is not None]


def _consume(values: list, n: Optional[int]) -> tuple:
    """Lấy n phần tử đầu (n=None -> lấy hết); trả (chunk, phần còn lại)."""
    if n is None:
        return values, []
    return values[:n], values[n:]


def _leftover_note(table_id: str, leftover: list) -> str:
    if not leftover:
        return ""
    return (f"Kịch bản đẩy dư {len(leftover)} giá trị report_val cho bảng "
            f"{table_id} không dùng tới — kiểm tra lại kịch bản.")


def apply_pass_rule(descriptor: TableDescriptor, row_def, raw_readings: list) -> tuple:
    """(measured, error, limit, passed) từ 1 chunk raw_readings theo đúng
    descriptor.pass_rule — dùng chung bởi map_table() lúc chạy kịch bản VÀ
    bởi gui/report_preview.py::_build_generic khi kiểm định viên sửa tay 1
    giá trị đo (double-click), đảm bảo công thức tính lại luôn NHẤT QUÁN
    với lúc map ban đầu."""
    rtype = descriptor.pass_rule.get("type", "none")
    # measured_count = số phần tử ĐẦU thật sự đại diện giá trị đo (phần còn
    # lại, nếu có, là field kịch bản tự tính sẵn — vd "sai số" — chỉ để
    # HIỂN THỊ lại trong Biên Bản qua report_val(), không dùng lại ở đây,
    # tránh trộn 2 quantity khác nhau vào 1 phép trung bình sai nghĩa.
    n_measured = row_def.measured_count if row_def.measured_count is not None else len(raw_readings)
    measured_readings = raw_readings[:n_measured]
    if len(measured_readings) != 1:
        measured = sum(measured_readings) / len(measured_readings) if measured_readings else None
    else:
        measured = measured_readings[0]

    error = None
    passed = None
    limit = row_def.limit

    if rtype == "relative_error_vs_fixed_limit":
        params = descriptor.pass_rule["params"]
        limit = params["limit_str"]
        error, passed = _relative_error_generic(measured, row_def.reference, params["fixed_limit"])
    elif rtype == "value_vs_parsed_threshold":
        passed = recompute_sensitivity(measured, row_def.limit)
    elif rtype == "correction_vs_reference":
        error = (row_def.reference - measured) if measured is not None else None
        passed = None
    # rtype == "none": measured/error/passed giữ nguyên (None/None)

    ui = row_def.uncertainty_index
    if ui is not None and ui < len(raw_readings):
        u_fmt = (row_def.value_format_seq[ui] if row_def.value_format_seq and ui < len(row_def.value_format_seq)
                 else descriptor.value_format)
        limit = _format(u_fmt, raw_readings[ui])

    return measured, error, limit, passed


def recompute_row(descriptor: TableDescriptor, row_index: int, raw_readings: list) -> Optional[tuple]:
    """Tính lại (measured, error, limit, passed) cho ĐÚNG 1 dòng (theo thứ
    tự trong descriptor.rows, khớp 1-1 với thứ tự ReportTable.rows trả về
    từ map_table()) sau khi kiểm định viên sửa tay 1/nhiều report_val của
    dòng đó. Trả None nếu row_index không hợp lệ."""
    if row_index < 0 or row_index >= len(descriptor.rows):
        return None
    return apply_pass_rule(descriptor, descriptor.rows[row_index], raw_readings)


def map_table(descriptor: TableDescriptor, step_results) -> ReportTable:
    """Thay cho mọi map_table_aN viết tay — duyệt descriptor.rows, tiêu thụ
    đúng số report_val (raw_count) cho mỗi dòng, áp công thức theo
    descriptor.pass_rule."""
    values = _table_values(step_results)
    rows = []
    table_passed = True
    rtype = descriptor.pass_rule.get("type", "none")

    for row_def in descriptor.rows:
        chunk, values = _consume(values, row_def.raw_count)
        # raw_readings LUÔN lưu nguyên chunk đã tiêu thụ cho dòng này (kể cả
        # raw_count=1) — đây là nguồn DUY NHẤT cho report_val() cursor
        # (build_cursor_context) đọc tuần tự, phải khớp CHÍNH XÁC những gì
        # kịch bản đã đẩy, không phải giá trị trung bình/suy diễn.
        raw_readings = chunk
        measured, error, limit, passed = apply_pass_rule(descriptor, row_def, raw_readings)

        if rtype == "relative_error_vs_fixed_limit" and error is not None and not passed:
            table_passed = False
        elif rtype == "value_vs_parsed_threshold" and passed is False:
            table_passed = False

        rows.append(TableRow(
            key=row_def.display_label or row_def.key,
            freq_set=row_def.freq_set,
            value_measured=measured,
            value_unit=descriptor.value_unit,
            error=error,
            limit=limit,
            passed=passed,
            raw_readings=raw_readings,
        ))

    if rtype in ("relative_error_vs_fixed_limit", "value_vs_parsed_threshold"):
        final_passed = table_passed if rows else None
    else:
        final_passed = None

    return ReportTable(
        table_id=descriptor.table_id,
        name=descriptor.name,
        rows=rows,
        passed=final_passed,
        note=_leftover_note(descriptor.table_id, values),
    )


# ---------------------------------------------------------------------------
# Format catalog — tên format (string, lưu trong descriptor JSON) -> hàm
# _fmt_*/_sci/... đã có sẵn trong report_generator.py/report_generator_nrp2.py.
# Chỉ thêm bảng tra, KHÔNG viết lại công thức định dạng.
# ---------------------------------------------------------------------------

_FORMATTERS = {
    "freq": _fmt_freq,
    "hz_measured": _fmt_hz_measured,
    "period": _fmt_period,
    "mv": _fmt_mv,
    "dbm": _fmt_dbm,
    "w": _fmt_w,
    "sci": _sci,
    "sci_signed": lambda v: f"± {_sci(v)}",
    "correction_mw": lambda v: _fmt_correction(v, "mW"),
    "correction_db": lambda v: _fmt_correction(v, "dB"),
    "text": lambda v: v if v is not None else "",
    # Biến thể "không kèm đơn vị" — cùng công thức làm tròn/định dạng số ở
    # trên, chỉ bỏ chữ đơn vị suffix, dùng khi cần bind giá trị thuần vào
    # docx (đơn vị đã có sẵn trong text mẫu hoặc không cần hiện).
    "freq_no_unit": lambda v: _fmt_freq(v, with_unit=False),
    "hz_measured_no_unit": lambda v: _fmt_hz_measured(v, with_unit=False),
    "period_no_unit": lambda v: _fmt_period(v, with_unit=False),
    "mv_no_unit": lambda v: _fmt_mv(v, with_unit=False),
    "dbm_no_unit": lambda v: _fmt_dbm(v, with_unit=False),
    "w_no_unit": lambda v: _fmt_w(v, with_unit=False),
    "correction_mw_no_unit": lambda v: _fmt_correction(v, "mW", with_unit=False),
    "correction_db_no_unit": lambda v: _fmt_correction(v, "dB", with_unit=False),
}


def _format(fmt: str, value) -> str:
    if value is None:
        return ""
    return _FORMATTERS[fmt](value)


def _extract_value(row: TableRow, role: str):
    if role == "freq_set":
        return row.freq_set
    if role == "measured":
        return row.value_measured
    if role == "error":
        return row.error
    if role == "limit":
        return row.limit
    if role == "free_text":
        return row.key
    if role == "power_from_key":
        return _power_set_from_key(row.key)
    raise ValueError(f"role '{role}' không dùng để trích 1 giá trị đơn (raw_reading_index xử lý riêng)")


# ---------------------------------------------------------------------------
# Generic context builder — list[TableRow] -> dict Jinja
# ---------------------------------------------------------------------------

def _spread_raw_readings(target: dict, row: TableRow, column, n: int):
    for i in range(n):
        v = row.raw_readings[i] if i < len(row.raw_readings) else None
        target[f"{column.jinja_field}{i + 1}"] = _format(column.format, v)


def _build_context_repeated_rows(descriptor: TableDescriptor, rows: list) -> dict:
    row_cols = [c for c in descriptor.columns if c.scope == "row"]
    table_cols = [c for c in descriptor.columns if c.scope == "table"]
    n_raw = descriptor.rows[0].raw_count if descriptor.rows else None

    ctx_rows = []
    for r in rows:
        d = {}
        for c in row_cols:
            if c.role == "raw_reading_index":
                _spread_raw_readings(d, r, c, n_raw or 0)
            else:
                d[c.jinja_field] = _format(c.format, _extract_value(r, c.role))
        ctx_rows.append(d)

    ctx = {"enabled": bool(rows), "rows": ctx_rows}
    for c in table_cols:
        if c.role == "limit" and descriptor.pass_rule.get("type") == "relative_error_vs_fixed_limit":
            ctx[c.jinja_field] = descriptor.pass_rule["params"]["limit_str"]
        else:
            v = _extract_value(rows[0], c.role) if rows else None
            ctx[c.jinja_field] = _format(c.format, v)
    return ctx


def _build_context_raw_expand_vertical(descriptor: TableDescriptor, rows: list) -> dict:
    """A1 CNT90XL — 1 dòng logic duy nhất, nhiều lần đo trải DỌC (docxtpl
    for-loop trên list "raws"), các cột còn lại là hằng số toàn bảng."""
    row0 = rows[0] if rows else None
    raw_col = next(c for c in descriptor.columns if c.role == "raw_reading_index")
    ctx = {"enabled": row0 is not None}
    if row0:
        for i, v in enumerate(row0.raw_readings):
            ctx.setdefault("raws", []).append({"idx": i + 1, raw_col.jinja_field: _format(raw_col.format, v)})
    else:
        ctx["raws"] = []

    for c in descriptor.columns:
        if c.role == "raw_reading_index":
            continue
        if c.role == "limit" and descriptor.pass_rule.get("type") == "relative_error_vs_fixed_limit":
            ctx[c.jinja_field] = descriptor.pass_rule["params"]["limit_str"]
        elif c.role == "freq_set" and row0 is None and descriptor.rows:
            # Bảng đang tắt/chưa có dữ liệu -> vẫn hiện tần số DANH ĐỊNH của
            # bảng (vd A1 luôn "10 MHz") thay vì để trống, giữ đúng hành vi cũ.
            ctx[c.jinja_field] = _format(c.format, descriptor.rows[0].freq_set)
        else:
            v = _extract_value(row0, c.role) if row0 else None
            ctx[c.jinja_field] = _format(c.format, v)
    return ctx


def _build_context_raw_expand_horizontal(descriptor: TableDescriptor, rows: list) -> dict:
    """NRP2 A1 — 1 dòng logic duy nhất, N lần đo trải NGANG thành N cột cố
    định (raw1..rawN), không cần vòng lặp for trong template."""
    row0 = rows[0] if rows else None
    n_raw = descriptor.rows[0].raw_count if descriptor.rows else 0
    ctx = {"enabled": row0 is not None}
    for c in descriptor.columns:
        if c.role == "raw_reading_index":
            if row0:
                _spread_raw_readings(ctx, row0, c, n_raw)
            else:
                for i in range(n_raw):
                    ctx[f"{c.jinja_field}{i + 1}"] = ""
        else:
            v = _extract_value(row0, c.role) if row0 else None
            if c.role == "limit" and descriptor.pass_rule.get("type") == "relative_error_vs_fixed_limit":
                ctx[c.jinja_field] = descriptor.pass_rule["params"]["limit_str"]
            else:
                ctx[c.jinja_field] = _format(c.format, v)
    return ctx


def build_table_context(descriptor: TableDescriptor, rows: list) -> dict:
    if descriptor.layout == "raw_expand_vertical":
        return _build_context_raw_expand_vertical(descriptor, rows)
    if descriptor.layout == "raw_expand_horizontal":
        return _build_context_raw_expand_horizontal(descriptor, rows)
    return _build_context_repeated_rows(descriptor, rows)


# ---------------------------------------------------------------------------
# Cursor context — cho phép quản trị viên tự gõ tag Jinja tuần tự thẳng vào
# bảng gốc trong Word (KHÔNG cần dựng {%tr for%}/vòng lặp) — mỗi lần gọi 1
# hàm trả về phần tử TIẾP THEO, đúng thứ tự tài liệu (trái→phải, trên→dưới,
# vì docxtpl render Jinja bằng cách duyệt 1 lượt cây XML tuần tự). Bổ sung
# THÊM vào context đã có (build_table_context) — KHÔNG thay field cũ, an
# toàn với 2 template cũ CNT90XL/NRP2 (chúng không tham chiếu field mới).
# ---------------------------------------------------------------------------

def _cursor(values: list):
    it = iter(values)

    def _next():
        return next(it, "")
    return _next


# NGUYÊN TẮC CHO BIÊN BẢN (khách hàng chốt): CHỈ 1 mảng report_val() duy
# nhất — quản trị viên/kịch bản TỰ tính toán MỌI giá trị cần hiện rồi tự đẩy
# (report_val step) vào ĐÚNG vị trí, ĐÚNG thứ tự đọc tài liệu (trái→phải,
# trên→dưới) — phần mềm KHÔNG được tự suy ra thêm field nào khác trong BIÊN
# BẢN (bản ghi thô, phải giữ nguyên những gì đã đo). report_val() đọc TUẦN
# TỰ từ `raw_readings` của từng dòng đã xác nhận, gộp phẳng đúng thứ tự dòng
# — raw_readings LÀ đúng những gì kịch bản đã đẩy (map_table() luôn lưu
# nguyên), KHÔNG phải giá trị trung bình/suy diễn.
#
# NGOẠI LỆ CHO GCN (Giấy chứng nhận): GCN là văn bản TỔNG HỢP/KẾT LUẬN, không
# phải bản ghi thô — được phép dùng field phần mềm tổng hợp, với điều kiện
# BẮT BUỘC: chỉ tổng hợp lại từ dữ liệu report_val Biên Bản ĐÃ CÓ SẴN (qua
# map_table()/pass_rule vẫn tính như cũ), KHÔNG YÊU CẦU sửa/thêm gì vào kịch
# bản đang chạy (khách hàng không thể sửa hàng loạt kịch bản đã viết xong).
#   - `result`      — Đạt/Không đạt CHÍNH NGƯỜI DÙNG tự chọn tay ở cột Bước 2
#                     (gui/report_preview.py::_add_status_column), không
#                     phải công thức tự động — "người dùng tự quyết định".
#                     NGOẠI LỆ: nếu có đúng 1 dòng đã xác nhận được đánh dấu
#                     "Xuất value trong GCN" (click phải 1 ô giá trị đo, xem
#                     gui/report_preview.py::_make_gcn_markable), `result`
#                     trả về giá trị đo ĐÃ ĐỊNH DẠNG của ô đó thay vì Đạt/
#                     Không đạt (_gcn_export_value_str) — mặc định (không
#                     dòng nào đánh dấu) vẫn là Đạt/Không đạt như cũ.
#   - `gcn_avg()`   — trung bình các report_val Biên Bản đã đẩy cho 1 dòng
#                     (TableRow.value_measured, map_table() đã tính sẵn).
#   - `gcn_error()` — sai số/số hiệu chỉnh tính từ report_val Biên Bản so
#                     với reference khai báo (TableRow.error, theo đúng công
#                     thức pass_rule của bảng — đã tính sẵn, không đổi).
#   - `gcn_limit()` — ngưỡng/giới hạn khai báo cho dòng đó (TableRow.limit).
# 3 cursor GCN này CHỈ NÊN dùng trong file gcnkd.docx — Biên Bản chỉ dùng
# report_val().
_ERROR_FORMAT_BY_VALUE_FORMAT = {"dbm": "correction_db", "w": "correction_mw"}


def _report_val_values(descriptor: TableDescriptor, rows: list) -> list:
    """Danh sách ĐÃ ĐỊNH DẠNG cho report_val(), tuần tự theo đúng thứ tự
    kịch bản đã đẩy. Mỗi dòng dùng đúng format của MÌNH — row_def.value_format_seq
    nếu bảng có dòng cần nhiều field/dòng với định dạng khác nhau (vd A5:
    giá trị đo rồi đến sai số kịch bản tự tính sẵn), mặc định là
    descriptor.value_format lặp lại cho MỌI report_val() (hành vi cũ)."""
    fmt = descriptor.value_format
    out = []
    for row_def, r in zip(descriptor.rows, rows):
        formats = row_def.value_format_seq or [fmt] * len(r.raw_readings)
        for v, f in zip(r.raw_readings, formats):
            out.append(_format(f, v))
    return out


def _gcn_export_value_str(descriptor: TableDescriptor, rt: Optional[ReportTable]) -> Optional[str]:
    """Nếu có ĐÚNG 1 dòng đã xác nhận trong `rt` được đánh dấu "Xuất value
    trong GCN" (TableRow.gcn_export_field = "raw:<idx>") -> trả giá trị đó
    đã định dạng ĐÚNG NHƯ report_val() sẽ hiện trong Biên Bản (dùng
    row_def.value_format_seq nếu dòng có, không thì descriptor.value_format).
    None nếu không có dòng nào đánh dấu -> `result` giữ hành vi cũ (Đạt/
    Không đạt). Khớp chỉ số rt.rows[i] <-> descriptor.rows[i] theo đúng bất
    biến của map_table() (1 TableRow/1 RowDef, cùng thứ tự, không lọc)."""
    if rt is None:
        return None
    for i, row in enumerate(rt.rows):
        if not row.confirmed or not row.gcn_export_field:
            continue
        if not row.gcn_export_field.startswith("raw:"):
            continue
        try:
            idx = int(row.gcn_export_field[len("raw:"):])
        except ValueError:
            continue
        if idx < 0 or idx >= len(row.raw_readings):
            continue
        fmt = descriptor.value_format
        if i < len(descriptor.rows):
            row_def = descriptor.rows[i]
            if row_def.value_format_seq and idx < len(row_def.value_format_seq):
                fmt = row_def.value_format_seq[idx]
        return _format(fmt, row.raw_readings[idx])
    return None


def build_cursor_context(descriptor: TableDescriptor, rows: list,
                          rt: Optional[ReportTable]) -> dict:
    fmt = descriptor.value_format
    error_fmt = _ERROR_FORMAT_BY_VALUE_FORMAT.get(fmt, "sci")
    export_value = _gcn_export_value_str(descriptor, rt)
    result = export_value if export_value is not None else (_pass_mark(rt.confirmed_passed) if rt else "")
    return {
        # _format() chỉ ĐỊNH DẠNG HIỂN THỊ (số thực -> chuỗi kiểu Việt Nam,
        # vd 12.3 -> "12,3 mVrms") — KHÔNG tính ra giá trị mới, giá trị vẫn
        # y hệt những gì kịch bản đã đẩy.
        "report_val": _cursor(_report_val_values(descriptor, rows)),
        "result": result,
        "gcn_avg": _cursor([_format(fmt, r.value_measured) for r in rows]),
        "gcn_error": _cursor([_format(error_fmt, r.error) for r in rows]),
        "gcn_limit": _cursor([r.limit or "" for r in rows]),
    }


# ---------------------------------------------------------------------------
# Marker ẩn + gộp ô hậu-render
# ---------------------------------------------------------------------------

_MARKER_RE = re.compile(r"##TABLE:([A-Za-z0-9_]+)##")
_MARKER_FMT = "##TABLE:{table_id}##"


def mark_table(tbl, table_id: str) -> None:
    """Chèn marker ẩn vào cell(0,0) của dòng header — gọi 1 lần lúc DỰNG file
    mẫu (scripts/build_seed_templates.py), KHÔNG gọi lúc render thật."""
    cell = tbl.cell(0, 0)
    p = cell.add_paragraph()
    run = p.add_run(_MARKER_FMT.format(table_id=table_id))
    run.font.size = Pt(1)
    run.font.hidden = True


def scan_table_markers(doc) -> dict:
    """Dò mọi bảng trong doc đã render, khớp marker ẩn -> {table_id: Table}.
    Bảng không có marker (header nhà nước, bảng chữ ký...) không nằm trong
    kết quả — thay hẳn cơ chế dò theo VỊ TRÍ cũ."""
    markers = {}
    for tbl in doc.tables:
        if not tbl.rows or not tbl.columns:
            continue
        m = _MARKER_RE.search(tbl.cell(0, 0).text)
        if m:
            markers[m.group(1)] = tbl
    return markers


def _merge_col_clean(tbl, col: int, row_start: int, row_end: int, text: str) -> None:
    """Gộp ô rồi dọn sạch — mỗi ô lặp dòng docxtpl đã có SẴN chữ (lặp lại y
    hệt nhau qua từng dòng), nên cell.merge() của python-docx sẽ NỐI CHUỖI
    các đoạn văn của từng ô lại (vd "10 MHz 10 MHz 10 MHz") thay vì thay thế.
    Phải xoá các đoạn văn thừa rồi ghi lại đúng 1 lần."""
    _merge_col(tbl, col, row_start, row_end)
    cell = tbl.cell(row_start, col)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    _cell_para(cell, text, size=SIZE_SMALL)


def _merge_grouped_col(tbl, rows: list, col: int, key_field: str, text_field: str) -> None:
    """Gộp các dòng LIÊN TIẾP có cùng giá trị `key_field` lại thành 1 ô ở
    cột `col`, hiển thị `text_field` của dòng đầu nhóm."""
    if not rows:
        return
    start = 0
    for i in range(1, len(rows) + 1):
        if i == len(rows) or rows[i][key_field] != rows[start][key_field]:
            if i - start > 1:
                _merge_col_clean(tbl, col, 1 + start, i, rows[start][text_field])   # +1: bỏ qua dòng header
            start = i


def apply_merges(tbl, ctx: dict, descriptor: TableDescriptor) -> None:
    if descriptor.layout == "raw_expand_vertical":
        n = len(ctx.get("raws", []))
    else:
        n = len(ctx.get("rows", []))

    for m in descriptor.merge:
        if m.mode == "constant":
            if n > 1:
                _merge_col_clean(tbl, m.col, 1, n, ctx[m.value_field])
        elif m.mode == "grouped":
            _merge_grouped_col(tbl, ctx["rows"], m.col, m.key_field, m.text_field)


# ---------------------------------------------------------------------------
# Generic renderer
# ---------------------------------------------------------------------------

def _find_report_table(session: CalibrationSession, table_id: str) -> Optional[ReportTable]:
    for test in session.tests:
        if test.table_id == table_id and test.result_table:
            return test.result_table
    return None


def build_all_table_contexts(session: CalibrationSession, descriptors: list) -> dict:
    """dict[table_id] -> context, dùng chung cho cả Biên Bản và GCN (NRP2) —
    2 file mẫu chỉ tham chiếu tới field jinja_field nào chúng cần, field dư
    không dùng tới bị Jinja bỏ qua, không cần 2 hàm build context riêng."""
    enabled_ids = {t.table_id for t in session.tests if t.enabled}
    result = {}
    for d in descriptors:
        rt = _find_report_table(session, d.table_id) if d.table_id in enabled_ids else None
        rows = rt.confirmed_rows() if rt else []
        ctx = build_table_context(d, rows)
        ctx.update(build_cursor_context(d, rows, rt))
        result[d.table_id] = ctx
    return result


def render_with_table_contexts(session: CalibrationSession, descriptors: list,
                                template_path, output_path, meta_context_fn) -> Path:
    """Render 1 file .docx có {"tables": {...}} trong context (Biên Bản
    CNT90XL/NRP2, GCN NRP2, và mọi mẫu quản trị viên tự gắn tag tay) — sau
    khi render, dò marker + gộp ô hậu-render theo descriptor.merge. Marker
    ẩn (`table_engine.mark_table`) CHỈ dùng để tìm bảng vật lý phục vụ gộp
    ô — bảng KHÔNG có `merge` (mọi bảng đăng ký qua luồng tự gõ tag tay,
    không dựng bảng/mutate gì nên không có marker) không cần marker, không
    lỗi. Bảng CÓ merge mà thiếu marker -> lỗi rõ ràng, dừng xuất (không âm
    thầm fallback về dò theo vị trí)."""
    output_path = Path(output_path)
    context = meta_context_fn(session)
    context["tables"] = build_all_table_contexts(session, descriptors)

    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    tpl.save(str(output_path))

    from docx import Document
    doc = Document(str(output_path))
    markers = scan_table_markers(doc)

    # markers rỗng -> tài liệu này KHÔNG dùng quy ước bảng-vật-lý/marker chút
    # nào (vd bảng tổng hợp GCN phẳng, mỗi bảng chỉ 1 dòng tĩnh + `tables.X.result`,
    # không có `{%tr for%}`/mark_table() nào) -> không có gì để gộp ô, bỏ qua
    # toàn bộ bước này thay vì báo lỗi. markers KHÔNG rỗng (tài liệu CÓ dùng
    # quy ước marker cho ít nhất 1 bảng) -> vẫn bắt buộc marker cho MỌI bảng
    # có merge như cũ, để bắt lỗi rõ ràng khi khách lỡ xoá dòng tiêu đề 1 bảng.
    if markers:
        for d in descriptors:
            if not d.merge:
                continue
            ctx = context["tables"][d.table_id]
            if not ctx.get("enabled"):
                continue
            if d.table_id not in markers:
                raise ValueError(
                    f"Không tìm thấy marker ẩn cho bảng {d.table_id} trong file mẫu "
                    f"'{template_path}' — có thể tiêu đề bảng đã bị ghi đè khi chỉnh sửa "
                    f"trong Word. Hãy mở lại từ file mẫu gốc để khôi phục bảng này."
                )
            apply_merges(markers[d.table_id], ctx, d)

    doc.save(str(output_path))
    return output_path


def render_gcnkd_summary(session: CalibrationSession, descriptors: list,
                          template_path, output_path, meta_context_fn) -> Path:
    """GCN CNT90XL — bảng tổng hợp 1 dòng/bài test (đạt/không đạt), KHÔNG
    dùng cấu trúc "tables"/descriptor.columns — chỉ dùng descriptor.gcn."""
    output_path = Path(output_path)
    enabled_ids = {t.table_id for t in session.tests if t.enabled}

    rows = []
    for d in sorted(descriptors, key=lambda x: x.order):
        if d.table_id not in enabled_ids or not d.gcn:
            continue
        rt = _find_report_table(session, d.table_id)
        result_str = _pass_mark(rt.confirmed_passed) if rt else ""
        rows.append({"param_name": d.gcn["param_name"], "result_str": result_str,
                      "limit": d.gcn["limit_str"]})

    context = meta_context_fn(session)
    context["rows"] = rows

    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    tpl.save(str(output_path))
    return output_path
