"""
core/report_generator_nrp2.py
==============================
Sinh tài liệu Word (.docx) theo mẫu QTHC 2.515 : 2021 (Máy đo công suất
NRP2) từ CalibrationSession — TÁCH RIÊNG khỏi core/report_generator.py
(mẫu CNT-90XL) vì đây là văn bản HIỆU CHUẨN (không có đạt/không đạt, không
có dòng Kết luận), khác cấu trúc với QTKĐ 2.461 (kiểm định).

Hai loại tài liệu:
  • generate_bienban() → Biên Bản Hiệu Chuẩn (Phụ lục A)
  • generate_gcnkd()   → Giấy Chứng Nhận Hiệu Chuẩn (Phụ lục B)

Tái dùng các hàm định dạng/docx thuần (không đổi) từ core/report_generator.py.
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from core.session import CalibrationSession, ReportTable, TableRow
from core.report_generator import (
    SIZE_NORMAL, SIZE_TITLE, SIZE_SMALL,
    _para, _cell_para, _set_col_widths, _merge_col, _clear_cell_borders,
    _add_run, _setup_page, _set_font, _add_gov_header, _sci,
    _fmt_freq, _fmt_dbm,
)

log_name = __name__


# ---------------------------------------------------------------------------
# Định dạng số riêng cho NRP2 (công suất W tại Bảng A1, số hiệu chỉnh)
# ---------------------------------------------------------------------------

def _fmt_w(w: float) -> str:
    """0.0010003 -> '1,0003 mW'."""
    mw = w * 1000.0
    return f"{mw:.4g} mW".replace(".", ",")


def _fmt_correction(value: Optional[float], unit: str) -> str:
    if value is None:
        return ""
    s = f"{value:.4g}".replace(".", ",")
    return f"{s} {unit}"


def _find_table(session: CalibrationSession, table_id: str) -> Optional[ReportTable]:
    for test in session.tests:
        if test.table_id == table_id and test.result_table:
            return test.result_table
    return None


def _confirmed(rt: Optional[ReportTable]) -> Optional[ReportTable]:
    if rt is None:
        return None
    return ReportTable(table_id=rt.table_id, name=rt.name,
                       rows=rt.confirmed_rows(), passed=rt.passed)


def _empty_table(table_id: str, name: str) -> ReportTable:
    return ReportTable(table_id=table_id, name=name, rows=[], passed=None)


# ---------------------------------------------------------------------------
# Bảng A1 — Độ chính xác mức công suất tại đầu ra chuẩn
# ---------------------------------------------------------------------------

_A1_N_LAN = 10   # mục 5.3.1: đo 10 lần lặp lại


def _add_table_a1_nrp2(doc: Document, rt: ReportTable):
    """Đúng mẫu trang 14: 1 dòng dữ liệu, cột 'lần 1'..'lần 10' nằm ngang
    (giống cách trải cột của Bảng A2/A3) — KHÔNG có cột TB/Số hiệu chỉnh
    trong Biên Bản (chỉ có ở GCN)."""
    _para(doc, "Bảng A1 - Xác định độ chính xác mức công suất tại đầu ra chuẩn",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_SMALL,
          space_before=4, space_after=2)

    row0 = rt.rows[0] if rt.rows else None
    raws = row0.raw_readings if row0 else []

    headers = ["Công suất\nchuẩn"] + [f"lần {i + 1}" for i in range(_A1_N_LAN)] + ["Độ\nKĐBĐ"]
    tbl = doc.add_table(rows=2, cols=len(headers))
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [2.2] + [1.0] * _A1_N_LAN + [1.8])

    for j, h in enumerate(headers):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    _cell_para(tbl.cell(1, 0), "1 mW", size=SIZE_SMALL)
    for i in range(_A1_N_LAN):
        if i < len(raws):
            _cell_para(tbl.cell(1, 1 + i), _fmt_w(raws[i]), size=SIZE_SMALL)
    _cell_para(tbl.cell(1, 1 + _A1_N_LAN), row0.limit if row0 else "", size=SIZE_SMALL)


# ---------------------------------------------------------------------------
# Bảng A2 — Độ chính xác đo mức công suất tuyệt đối (tại 0 dBm)
# ---------------------------------------------------------------------------

def _add_table_a2_nrp2(doc: Document, rt: ReportTable):
    """Đúng mẫu trang 15: Tần số | 5 lần đo + TB | Độ KĐBĐ — KHÔNG có cột
    Số hiệu chỉnh trong Biên Bản (chỉ có ở GCN)."""
    _para(doc, "Bảng A2 - Xác định độ chính xác đo mức công suất tuyệt đối (tại 0 dBm)",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_SMALL,
          space_before=4, space_after=2)

    n_rows = 1 + len(rt.rows)
    tbl = doc.add_table(rows=n_rows, cols=8)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [3.5, 1.6, 1.6, 1.6, 1.6, 1.6, 1.8, 2.5])

    headers = ["Tần số thiết lập\n(mức công suất 0 dBm)",
               "lần 1", "lần 2", "lần 3", "lần 4", "lần 5", "TB", "Độ KĐBĐ"]
    for j, h in enumerate(headers):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    for i, r in enumerate(rt.rows, start=1):
        if r.freq_set:
            _cell_para(tbl.cell(i, 0), _fmt_freq(r.freq_set), size=SIZE_SMALL)
        for k in range(5):
            if k < len(r.raw_readings):
                _cell_para(tbl.cell(i, 1 + k), _fmt_dbm(r.raw_readings[k]), size=SIZE_SMALL)
        if r.value_measured is not None:
            _cell_para(tbl.cell(i, 6), _fmt_dbm(r.value_measured), size=SIZE_SMALL)
        _cell_para(tbl.cell(i, 7), r.limit, size=SIZE_SMALL)


# ---------------------------------------------------------------------------
# Bảng A3 — Độ chính xác đo công suất với bộ NRPC50 calibration kit
# ---------------------------------------------------------------------------

def _add_table_a3_nrp2(doc: Document, rt: ReportTable):
    """Đúng mẫu trang 15-16: Tần số | Công suất chuẩn | 5 lần đo + TB | Độ
    KĐBĐ — KHÔNG có cột Số hiệu chỉnh trong Biên Bản (chỉ có ở GCN)."""
    _para(doc, "Bảng A3 - Xác định độ chính xác đo công suất với bộ hiệu chuẩn "
               "công suất NRPC50 calibration kit",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_SMALL,
          space_before=4, space_after=2)

    n_rows = 1 + len(rt.rows)
    tbl = doc.add_table(rows=n_rows, cols=9)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [2.5, 2.5, 1.6, 1.6, 1.6, 1.6, 1.6, 1.7, 2.3])

    headers = ["Tần số\nthiết lập", "Công suất\nchuẩn (dBm)",
               "lần 1", "lần 2", "lần 3", "lần 4", "lần 5", "TB", "Độ\nKĐBĐ"]
    for j, h in enumerate(headers):
        _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)

    # Gộp ô cột "Tần số thiết lập" theo nhóm liên tiếp cùng freq_set
    rows = rt.rows
    start = 0
    for i in range(1, len(rows) + 1):
        if i == len(rows) or rows[i].freq_set != rows[start].freq_set:
            if i - start > 1:
                _merge_col(tbl, 0, start + 1, i)
            _cell_para(tbl.cell(start + 1, 0), _fmt_freq(rows[start].freq_set), size=SIZE_SMALL)
            start = i

    for i, r in enumerate(rows, start=1):
        power_set = _power_set_from_key(r.key)
        if power_set is not None:
            _cell_para(tbl.cell(i, 1), _fmt_dbm(power_set), size=SIZE_SMALL)
        for k in range(5):
            if k < len(r.raw_readings):
                _cell_para(tbl.cell(i, 2 + k), _fmt_dbm(r.raw_readings[k]), size=SIZE_SMALL)
        if r.value_measured is not None:
            _cell_para(tbl.cell(i, 7), _fmt_dbm(r.value_measured), size=SIZE_SMALL)
        _cell_para(tbl.cell(i, 8), r.limit, size=SIZE_SMALL)


def _power_set_from_key(key: str) -> Optional[float]:
    if key and "_" in key:
        try:
            return float(key.rsplit("_", 1)[1].replace("dBm", ""))
        except ValueError:
            return None
    return None


# ---------------------------------------------------------------------------
# BIÊN BẢN HIỆU CHUẨN (Phụ lục A) — dùng lại header nhà nước 2 cột của CNT-90XL
# ---------------------------------------------------------------------------

def _add_hieuchuan_tables(doc: Document, session: CalibrationSession):
    enabled = {t.table_id for t in session.tests if t.enabled}

    if "A1" in enabled:
        _para(doc, "A.3.1 Xác định độ chính xác mức công suất tại đầu ra chuẩn",
              bold=True, space_before=4, space_after=2)
        _add_table_a1_nrp2(doc, _confirmed(_find_table(session, "A1"))
                          or _empty_table("A1", ""))

    if "A2" in enabled:
        _para(doc, "A.3.2 Xác định độ chính xác đo mức công suất tuyệt đối",
              bold=True, space_before=6, space_after=2)
        _add_table_a2_nrp2(doc, _confirmed(_find_table(session, "A2"))
                          or _empty_table("A2", ""))

    if "A3" in enabled:
        _para(doc, "A.3.3 Xác định độ chính xác đo công suất với bộ hiệu chuẩn "
                   "công suất NRPC50 calibration kit",
              bold=True, space_before=6, space_after=2)
        _add_table_a3_nrp2(doc, _confirmed(_find_table(session, "A3"))
                          or _empty_table("A3", ""))


def generate_bienban(session: CalibrationSession, output_path: str | Path) -> Path:
    """Sinh Biên Bản Hiệu Chuẩn theo mẫu QTHC 2.515 : 2021 (Phụ lục A)."""
    output_path = Path(output_path)
    doc = Document()
    _setup_page(doc)
    for p in doc.paragraphs:
        p.clear()

    meta = session.meta
    dut = meta.dut

    _add_gov_header(doc, session)
    _para(doc, "", space_after=4)

    _para(doc, "BIÊN BẢN HIỆU CHUẨN",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_TITLE,
          space_before=6, space_after=4)

    def _info_line(label: str, value: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _add_run(p, label); _add_run(p, value)

    def _info_line_2col(label1, val1, label2, val2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _add_run(p, label1); _add_run(p, val1)
        _add_run(p, "          ")
        _add_run(p, label2); _add_run(p, val2)

    _info_line("Tên phương tiện ĐL-TN: ", "Máy đo công suất kiểu NRP2")
    _info_line_2col("Ký hiệu: ", dut.model or "NRP2", "Số hiệu: ", dut.serial)
    _info_line("Nước (Hãng) sản xuất: ", dut.manufacturer or "R&S")
    _info_line("Đặc tính đo lường: ",
              dut.measurement_range or "Dải tần làm việc từ DC đến 110 GHz; "
                                        "dải đo công suất từ (-67 đến 45) dBm")
    _info_line("Đơn vị sử dụng: ", dut.owner)
    _info_line("Phương pháp hiệu chuẩn: ", "QTHC 2.515 : 2021")
    _info_line("Phương tiện hiệu chuẩn: ", meta.inspection_equipment)
    _info_line_2col("Điều kiện môi trường:  nhiệt độ ", meta.temperature,
                    "      độ ẩm ", meta.humidity)

    if meta.date:
        d = meta.date
        _info_line("Đã tiến hành hiệu chuẩn, ngày ",
                  f"{d.day:02d} tháng {d.month:02d} năm {d.year}")
    else:
        _info_line("Đã tiến hành hiệu chuẩn, ngày ", "        tháng         năm 20")

    _para(doc, "KẾT QUẢ HIỆU CHUẨN",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_NORMAL,
          space_before=6, space_after=4)

    _para(doc, "A.1 Kiểm tra bên ngoài", bold=True, space_before=2, space_after=1)
    _para(doc, "A.2 Kiểm tra kỹ thuật",  bold=True, space_before=2, space_after=1)
    _para(doc, "A.3 Kiểm tra đo lường",  bold=True, space_before=2, space_after=2)

    _add_hieuchuan_tables(doc, session)

    # Không có đoạn "Kết luận" — QTHC 2.515 là văn bản hiệu chuẩn, không có
    # khái niệm đạt/không đạt (khác QTKĐ 2.461).

    _para(doc, "", space_after=8)
    tbl_sign = doc.add_table(rows=1, cols=2)
    tbl_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl_sign.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)

    lc = tbl_sign.cell(0, 0)
    rc = tbl_sign.cell(0, 1)
    for cell in (lc, rc):
        cell.width = Cm(8)

    lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(lp, "Người soát lại", bold=True)
    lc.add_paragraph(meta.reviewer or "").alignment = WD_ALIGN_PARAGRAPH.CENTER

    rp = rc.paragraphs[0]; rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(rp, "Kiểm định viên", bold=True)
    rc.add_paragraph(meta.operator or "").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# GIẤY CHỨNG NHẬN HIỆU CHUẨN (Phụ lục B) — header song ngữ 1 khối riêng
# ---------------------------------------------------------------------------

def _add_header_gcn_nrp2(doc: Document):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p1, "CỤC TIÊU CHUẨN ĐO LƯỜNG–CHẤT LƯỢNG\n", bold=True)
    _add_run(p1, "(Department for Standard, Metrology and Quality)\n", italic=True, size=SIZE_SMALL)
    _add_run(p1, "TRUNG TÂM TIÊU CHUẨN ĐO LƯỜNG–CHẤT LƯỢNG 2\n", bold=True)
    _add_run(p1, "(Standard, Metrology and Quality Centrer 2)\n", italic=True, size=SIZE_SMALL)
    _add_run(p1, "Địa chỉ (Add): Số 40 Nguyễn Giản Thanh – Phường 15 – Quận10 - TP. Hồ Chí Minh\n",
             size=SIZE_SMALL)
    _add_run(p1, "Điện thoại (Tel): 08.39700977 – Fax: 08.39700977", size=SIZE_SMALL)


def generate_gcnkd(session: CalibrationSession, output_path: str | Path) -> Path:
    """Sinh Giấy Chứng Nhận Hiệu Chuẩn theo mẫu QTHC 2.515 : 2021 (Phụ lục B)."""
    output_path = Path(output_path)
    doc = Document()
    _setup_page(doc)
    for p in doc.paragraphs:
        p.clear()

    meta = session.meta
    dut = meta.dut

    _add_header_gcn_nrp2(doc)
    _para(doc, "", space_after=6)

    _para(doc, "GIẤY CHỨNG NHẬN HIỆU CHUẨN",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=SIZE_TITLE,
          space_before=4, space_after=1)
    _para(doc, "(Calibration Certificate)",
          align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=SIZE_SMALL, space_after=2)
    _para(doc, f"Số (N0): {meta.cert_number}",
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    def _line2(lbl1, v1, lbl2, v2):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _add_run(p, lbl1); _add_run(p, v1)
        _add_run(p, "     ")
        _add_run(p, lbl2); _add_run(p, v2)

    def _line(lbl, val):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _add_run(p, lbl); _add_run(p, val)

    _line("Tên phương tiện đo (Object): ", "Máy đo công suất kiểu NRP2")
    _line2("Kiểu (Type): ", dut.model or "NRP2", "Số (Serial N0): ", dut.serial)
    _line("Nơi sản xuất (Manufacturer): ", dut.manufacturer or "R&S")
    _line("Đặc trưng kỹ thuật (Technical Specification): ",
        dut.measurement_range or "Dải tần làm việc từ DC đến 110 GHz; "
                                  "dải đo công suất từ (-67 đến 45) dBm")
    _line("Cơ sở sử dụng (Customer): ", dut.owner)
    _line("Phương pháp thực hiện (Method of calibration): ", "QTHC 2.515 : 2021")
    _line2("Điều kiện môi trường (Environmental Conditions):  nhiệt độ ", meta.temperature,
          "độ ẩm ", meta.humidity)
    _line("Chuẩn được sử dụng (Standards used): ", meta.inspection_equipment)

    if meta.date:
        d = meta.date
        _line("Ngày hiệu chuẩn (Date of Calibration): ",
             f"{d.day:02d} tháng {d.month:02d} năm {d.year}")

    _para(doc, "KẾT QUẢ HIỆU CHUẨN",
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_before=6, space_after=2)
    _para(doc, "(Calibration results)",
          align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=SIZE_SMALL, space_after=4)

    _add_gcn_result_tables(doc, session)

    _para(doc, "", space_after=4)
    _para(doc, "Ghi chú :", italic=True, size=SIZE_SMALL)
    _para(doc, "1. Độ không đảm bảo đo mở rộng được tính từ độ không đảm bảo đo chuẩn "
              "nhân với hệ số phủ k=2, phân bố chuẩn tương ứng với 95 % độ tin cậy.",
          size=SIZE_SMALL)
    _para(doc, "2. Các điểm hiệu chuẩn thực hiện theo yêu cầu của khách hàng.",
          size=SIZE_SMALL, space_after=6)

    if meta.location and meta.date:
        d = meta.date
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(p_date, f"{meta.location}, ngày {d.day:02d} tháng {d.month:02d} năm {d.year}",
                italic=True)

    tbl_sign = doc.add_table(rows=1, cols=2)
    tbl_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in tbl_sign.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)
            cell.width = Cm(8)

    titles = ["TRƯỞNG PHÒNG THÍ NGHIỆM\n(Head of the Cal. Lab.)",
              "GIÁM ĐỐC\n(Director)"]
    for j, title in enumerate(titles):
        cell = tbl_sign.cell(0, j)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(cp, title, bold=True, size=SIZE_SMALL)

    doc.save(str(output_path))
    return output_path


def _add_gcn_result_tables(doc: Document, session: CalibrationSession):
    """Bảng kết quả rút gọn trong GCN — chỉ TB + Số hiệu chỉnh + Độ KĐBĐ,
    không có cột từng lần đo riêng lẻ (khác Biên Bản)."""
    enabled = {t.table_id for t in session.tests if t.enabled}

    if "A1" in enabled:
        _para(doc, "1. Xác định độ chính xác mức công suất tại đầu ra chuẩn",
              bold=True, space_before=2, space_after=2)
        rt = _confirmed(_find_table(session, "A1")) or _empty_table("A1", "")
        tbl = doc.add_table(rows=1 + max(len(rt.rows), 1), cols=4)
        tbl.style = "Table Grid"
        _set_col_widths(tbl, [4.0, 4.5, 3.0, 3.0])
        for j, h in enumerate(["Công suất chuẩn\n(tại f = 50 MHz)",
                               "Công suất đo được\ntrên NRVD (tại f = 50 MHz)",
                               "Số hiệu chỉnh", "Độ KĐBĐ"]):
            _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)
        if rt.rows:
            r = rt.rows[0]
            _cell_para(tbl.cell(1, 0), "1 mW", size=SIZE_SMALL)
            if r.value_measured is not None:
                _cell_para(tbl.cell(1, 1), _fmt_w(r.value_measured), size=SIZE_SMALL)
            _cell_para(tbl.cell(1, 2), _fmt_correction(r.error, "mW"), size=SIZE_SMALL)
            _cell_para(tbl.cell(1, 3), r.limit, size=SIZE_SMALL)
        else:
            _cell_para(tbl.cell(1, 0), "1 mW", size=SIZE_SMALL)

    if "A2" in enabled:
        _para(doc, "2. Xác định độ chính xác đo mức công suất tuyệt đối (tại 0 dBm)",
              bold=True, space_before=6, space_after=2)
        rt = _confirmed(_find_table(session, "A2")) or _empty_table("A2", "")
        n = max(len(rt.rows), 1)
        tbl = doc.add_table(rows=1 + n, cols=4)
        tbl.style = "Table Grid"
        _set_col_widths(tbl, [4.0, 4.5, 3.0, 3.0])
        for j, h in enumerate(["Tần số thiết lập\n(mức công suất 0 dBm)",
                               "Công suất đo\nđược trên NRP2", "Số hiệu chỉnh", "Độ KĐBĐ"]):
            _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)
        for i, r in enumerate(rt.rows, start=1):
            if r.freq_set:
                _cell_para(tbl.cell(i, 0), _fmt_freq(r.freq_set), size=SIZE_SMALL)
            if r.value_measured is not None:
                _cell_para(tbl.cell(i, 1), _fmt_dbm(r.value_measured), size=SIZE_SMALL)
            _cell_para(tbl.cell(i, 2), _fmt_correction(r.error, "dB"), size=SIZE_SMALL)
            _cell_para(tbl.cell(i, 3), r.limit, size=SIZE_SMALL)

    if "A3" in enabled:
        _para(doc, "3. Xác định độ chính xác đo công suất",
              bold=True, space_before=6, space_after=2)
        rt = _confirmed(_find_table(session, "A3")) or _empty_table("A3", "")
        n = max(len(rt.rows), 1)
        tbl = doc.add_table(rows=1 + n, cols=5)
        tbl.style = "Table Grid"
        _set_col_widths(tbl, [3.0, 3.0, 4.5, 3.0, 3.0])
        for j, h in enumerate(["Tần số\nthiết lập", "Công suất\nchuẩn trên NRP2 (dBm)",
                               "Công suất trung bình\nđo được trên NRP2, dBm",
                               "Số hiệu\nchỉnh", "Độ\nKĐBĐ"]):
            _cell_para(tbl.cell(0, j), h, bold=True, size=SIZE_SMALL)
        rows = rt.rows
        start = 0
        for i in range(1, len(rows) + 1):
            if i == len(rows) or rows[i].freq_set != rows[start].freq_set:
                if i - start > 1:
                    _merge_col(tbl, 0, start + 1, i)
                if rows:
                    _cell_para(tbl.cell(start + 1, 0), _fmt_freq(rows[start].freq_set), size=SIZE_SMALL)
                start = i
        for i, r in enumerate(rows, start=1):
            power_set = _power_set_from_key(r.key)
            if power_set is not None:
                _cell_para(tbl.cell(i, 1), _fmt_dbm(power_set), size=SIZE_SMALL)
            if r.value_measured is not None:
                _cell_para(tbl.cell(i, 2), _fmt_dbm(r.value_measured), size=SIZE_SMALL)
            _cell_para(tbl.cell(i, 3), _fmt_correction(r.error, "dB"), size=SIZE_SMALL)
            _cell_para(tbl.cell(i, 4), r.limit, size=SIZE_SMALL)
