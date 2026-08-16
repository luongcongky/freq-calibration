"""
core/table_wizard_io.py
=========================
Phần "không Qt" của "Thêm bảng/mẫu báo cáo mới" (gui/template_manager_dialog.py
là lớp giao diện gọi vào đây, core/table_import.py điều phối ghi) — dựng
TableDescriptor (core/table_descriptor.py) từ dữ liệu người dùng nhập ở màn
hình review, ghi ra file JSON.

Quản trị viên TỰ TAY gõ tag Jinja trực tiếp vào file .docx thật trong Word
(xem cheat-sheet ở gui/template_manager_dialog.py) — file này KHÔNG còn đọc/
sửa cấu trúc bảng vật lý nào trong .docx nữa (khác bản trước có cơ chế
"row-surgery" tự động chèn tag).

Không phụ thuộc Qt → test độc lập với phần GUI.

Giới hạn đã biết (thống nhất khi thiết kế):
  - KHÔNG tự parse ngược chuỗi đã định dạng (vd "100 kHz" -> 100000.0) —
    guess_bare_number() chỉ nhận SỐ THUẦN, còn lại để người dùng xác nhận
    tay qua màn hình review.
"""

from __future__ import annotations

import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from core.table_descriptor import RowDef, TableDescriptor

# ---------------------------------------------------------------------------
# Model trung gian thu thập qua màn hình review
# ---------------------------------------------------------------------------

def guess_bare_number(text: str) -> Optional[float]:
    """CHỈ trả về số khi text là SỐ THUẦN (không đơn vị/tiền tố) — tránh
    parse sai kiểu '100 kHz' -> 100.0. Dùng để gợi ý điền sẵn, không phải
    nguồn dữ liệu chính thức (người dùng luôn xác nhận/gõ tay)."""
    if text is None:
        return None
    s = text.strip().replace(",", ".")
    if not s:
        return None
    try:
        return float(s)
    except ValueError:
        return None


@dataclass
class WizardRowSpec:
    key: str = ""
    freq_set: Optional[float] = None
    reference: Optional[float] = None
    limit: str = ""
    display_label: str = ""


@dataclass
class WizardTableSpec:
    table_id: str
    name: str
    order: int
    value_unit: str
    value_format: str = "text"          # 1 trong FORMAT_LABELS_ALL — áp dụng cho
                                         # report_val()/report_raw() của CẢ bảng
    rows: list = field(default_factory=list)         # list[WizardRowSpec]
    pass_rule: dict = field(default_factory=lambda: {"type": "none"})
    gcn: Optional[dict] = None


def is_advanced_table(d: TableDescriptor) -> bool:
    """True nếu có dòng nào khai raw_count > 1 (dòng nhận NHIỀU report_val()
    liên tiếp, cần value_format_seq riêng cho từng report_val()) — vẫn sửa
    được qua gui/template_manager_dialog.py::TableFormDialog (nút "Sửa"),
    KHÔNG bắt buộc mọi dòng cùng 1 raw_count (1 dòng tổng hợp gộp ô có thể
    ít report_val() hơn dòng đo thường, xem raw_counts_for_measured_cols()).
    Phần mềm không tự tính gì từ các report_val() phụ này (core/table_engine.py::
    apply_pass_rule) — chỉ hiển thị lại nguyên văn."""
    return any((r.raw_count or 1) > 1 for r in d.rows)


def descriptor_to_spec(d: TableDescriptor) -> WizardTableSpec:
    """Chiều NGƯỢC của build_descriptor() — nạp lại 1 bảng đã có thành
    WizardTableSpec để điền sẵn form sửa. KHÔNG tự lấy raw_count/
    value_format_seq (bảng nâng cao) — gui/template_manager_dialog.py::
    _edit_table() tự đọc trực tiếp từ d.rows (mỗi dòng 1 raw_count riêng)
    và truyền riêng cho TableFormDialog(raw_counts=..., row_value_format_seqs=...)."""
    rows = [WizardRowSpec(key=r.key, freq_set=r.freq_set, reference=r.reference,
                           limit=r.limit, display_label=r.display_label)
            for r in d.rows]
    return WizardTableSpec(table_id=d.table_id, name=d.name, order=d.order,
                            value_unit=d.value_unit, value_format=d.value_format,
                            rows=rows, pass_rule=d.pass_rule, gcn=d.gcn)


def validate_table_id_available(tables_dir, table_id: str) -> Optional[str]:
    """Trả về thông báo lỗi (str) nếu table_id không hợp lệ/đã tồn tại, None
    nếu hợp lệ."""
    if not table_id or not table_id.strip():
        return "Mã bảng không được để trống."
    if not table_id.replace("_", "").isalnum():
        return "Mã bảng chỉ được chứa chữ/số/gạch dưới (vd 'A9')."
    tables_dir = Path(tables_dir)
    if (tables_dir / f"{table_id}.json").exists():
        return f"Bảng '{table_id}' đã tồn tại — hãy chọn mã khác."
    return None


def validate_rows(rows: list, pass_rule: dict) -> Optional[str]:
    """Trả về thông báo lỗi nếu dữ liệu từng dòng chưa đủ theo pass_rule đã
    chọn, None nếu ổn."""
    if not rows:
        return "Cần ít nhất 1 dòng dữ liệu."
    keys = [r.key.strip() for r in rows]
    if any(not k for k in keys):
        return "Mọi dòng phải có Khoá (không được để trống)."
    if len(set(keys)) != len(keys):
        return "Khoá của các dòng phải khác nhau (không trùng)."
    rtype = pass_rule.get("type", "none")
    if rtype in ("relative_error_vs_fixed_limit", "correction_vs_reference"):
        if any(r.reference is None for r in rows):
            return "Mọi dòng phải có giá trị 'Chuẩn dùng để tính' (số) theo quy tắc Đạt/Không đạt đã chọn."
    if rtype == "value_vs_parsed_threshold":
        if any(not r.limit.strip() for r in rows):
            return "Mọi dòng phải có 'Ngưỡng' (vd '≤ 15 mVrms') theo quy tắc Đạt/Không đạt đã chọn."
    return None


# ---------------------------------------------------------------------------
# Dựng TableDescriptor từ WizardTableSpec
# ---------------------------------------------------------------------------

def build_descriptor(spec: WizardTableSpec) -> TableDescriptor:
    """Không còn khái niệm cột vật lý/gộp ô (`columns`/`merge` luôn rỗng) —
    quản trị viên tự gõ tag `report_val()`/... trực tiếp trong Word, mỗi
    dòng khai báo luôn tiêu thụ ĐÚNG 1 report_val (`raw_count=1`)."""
    row_defs = [
        RowDef(key=r.key, freq_set=r.freq_set, reference=r.reference,
               raw_count=1, limit=r.limit, display_label=r.display_label)
        for r in spec.rows
    ]
    return TableDescriptor(
        schema_version=1, table_id=spec.table_id, name=spec.name, order=spec.order,
        scenario_file="", layout="repeated_rows", value_unit=spec.value_unit,
        value_format=spec.value_format, rows=row_defs, columns=[],
        pass_rule=spec.pass_rule, merge=[], gcn=spec.gcn,
    )


def write_descriptor_json(descriptor: TableDescriptor, tables_dir) -> Path:
    tables_dir = Path(tables_dir)
    tables_dir.mkdir(parents=True, exist_ok=True)
    out_path = tables_dir / f"{descriptor.table_id}.json"
    out_path.write_text(json.dumps(descriptor.to_dict(), ensure_ascii=False, indent=2),
                         encoding="utf-8")
    return out_path


# ---------------------------------------------------------------------------
# Đọc file .docx khách ĐÃ TỰ gắn tag Jinja tay — CHỈ ĐỌC, không sửa gì.
# ---------------------------------------------------------------------------

@dataclass
class DetectedMetaPair:
    """Kết quả tách "Nhãn: giá trị" ngoài bảng — CHỈ để gợi ý điền sẵn form
    meta/dut khi tạo mẫu mới, KHÔNG phải nguồn dữ liệu chính thức."""
    label: str
    value: str


_META_COLON_RE = re.compile(r"^(.{2,60}?)\s*[:：]\s*(.+)$")


def scan_meta_paragraphs(docx_path) -> list:
    """Best-effort tách đoạn văn "Nhãn: giá trị" NGOÀI bảng — CHỈ dùng để
    gợi ý điền sẵn form meta/dut khi tạo mẫu mới, KHÔNG phải nguồn dữ liệu
    chính thức (người dùng luôn xác nhận/sửa trước khi lưu)."""
    doc = Document(str(docx_path))
    result = []
    for el in doc.element.body:
        if el.tag != qn("w:p"):
            continue
        text = Paragraph(el, doc).text.strip()
        m = _META_COLON_RE.match(text)
        if m:
            result.append(DetectedMetaPair(label=m.group(1).strip(), value=m.group(2).strip()))
    return result


def find_missing_table_ids(docx_path, table_ids: list) -> list:
    """Đọc toàn bộ text (đoạn văn + mọi ô bảng) của file .docx đã gắn tag
    tay, trả về danh sách table_id KHÔNG tìm thấy chuỗi 'tables.<ID>' ở đâu
    trong file — cảnh báo sớm lỗi gõ nhầm mã bảng, KHÔNG sửa/chặn gì."""
    doc = Document(str(docx_path))
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    full_text = "\n".join(parts)
    return [tid for tid in table_ids if f"tables.{tid}" not in full_text]


# ---------------------------------------------------------------------------
# "Đọc bảng từ Word" — khách đã tự dựng SẴN bảng thật (đủ dòng/cột, đủ nhãn/
# ngưỡng tĩnh) trong file .docx của họ, chỉ CHỪA TRỐNG ô sẽ chứa giá trị đo.
# Khác hẳn cơ chế "row-surgery" đã bỏ (không tự dựng/sửa số dòng, số cột, gộp
# ô nào cả) — chỉ ĐỌC text ô có sẵn để suy ra dữ liệu dòng, và GHI ĐÈ đúng
# text của những ô khách đã đánh dấu là "giá trị đo" thành tag report_val(),
# không đụng gì khác trong file.
# ---------------------------------------------------------------------------

@dataclass
class DetectedDocxTable:
    index: int                # vị trí bảng trong doc.tables (0-based) — dùng lại khi ghi tag
    n_rows: int
    n_cols: int
    grid: list                # list[list[str]] — text từng ô, nguyên văn đọc từ Word
    already_tagged: bool      # True nếu trong bảng đã có chữ "report_val()" hoặc "{{ tables."


def _visible_row_cells(row) -> list:
    """`row.cells` (python-docx) LẶP LẠI cùng 1 ô nhiều lần cho mỗi cột nó
    gridSpan (gộp ngang) qua — hành vi MẶC ĐỊNH của python-docx (không phải
    bug), nhưng nếu dùng thẳng để dựng grid sẽ ra NHIỀU cột hơn số cột NHÌN
    THẤY thật trong Word (vd 1 ô "lần 5" gộp 2 cột sẽ bị đọc thành 2 cột
    "lần 5" liên tiếp). Hàm này chỉ giữ lần xuất hiện ĐẦU TIÊN của mỗi ô vật
    lý (so theo identity ô XML thật `_tc`) — khớp ĐÚNG số cột thật, dùng
    chung bởi scan_docx_tables() (đọc) và insert_report_val_tags() (ghi) để
    2 bên LUÔN đánh số cột giống nhau (không thì chọn cột N lúc đọc sẽ ghi
    tag nhầm sang cột khác lúc viết)."""
    cells = []
    prev_tc = None
    for cell in row.cells:
        if cell._tc is prev_tc:
            continue
        prev_tc = cell._tc
        cells.append(cell)
    return cells


def _visible_row_cells_with_pos(row) -> list:
    """Giống _visible_row_cells nhưng trả kèm (cell, start_pos, span) — vị
    trí lưới (0-based) mà mỗi ô BẮT ĐẦU chiếm VÀ số cột lưới nó chiếm —
    `row.cells` đã pad đúng bằng độ rộng lưới khai báo (tblGrid) nên chỉ số
    enumerate() CHÍNH LÀ vị trí lưới thật. Cần cả span (không chỉ vị trí bắt
    đầu) để so khớp ĐÚNG vai trò cột đã gán ở DÒNG TIÊU ĐỀ với cấu trúc ô
    THẬT của TỪNG dòng dữ liệu — 1 dòng tổng hợp (vd "Trung Bình") có thể
    gộp nhiều ô hẹp thành 1 ô rộng khác hẳn dòng tiêu đề, và ngay cả DÒNG
    TIÊU ĐỀ cũng có thể có ô gộp (vd "lần 5" chiếm 2 cột lưới) — nếu chỉ
    biết vị trí bắt đầu của ô header đó, phần cột lưới THỨ 2 nó cũng chiếm
    sẽ bị bỏ sót khi so khớp với ô của dòng dữ liệu khác bắt đầu đúng ngay
    tại đó."""
    cells = []
    prev_tc = None
    start = 0
    span = 0
    prev_cell = None
    for pos, cell in enumerate(row.cells):
        if cell._tc is prev_tc:
            span += 1
            continue
        if prev_tc is not None:
            cells.append((prev_cell, start, span))
        prev_tc = cell._tc
        prev_cell = cell
        start = pos
        span = 1
    if prev_tc is not None:
        cells.append((prev_cell, start, span))
    return cells


def scan_docx_tables(docx_path) -> list:
    """Đọc TOÀN BỘ bảng vật lý cấp cao nhất (doc.tables — không gồm bảng lồng
    trong ô) của 1 file .docx, trả về list[DetectedDocxTable] theo đúng thứ
    tự xuất hiện trong tài liệu. CHỈ ĐỌC, không sửa gì trong file."""
    doc = Document(str(docx_path))
    result = []
    for i, tbl in enumerate(doc.tables):
        grid = [[cell.text.strip() for cell in _visible_row_cells(row)] for row in tbl.rows]
        flat = " ".join(c for r in grid for c in r)
        already_tagged = "report_val()" in flat or "{{ tables." in flat
        result.append(DetectedDocxTable(
            index=i, n_rows=len(grid), n_cols=(len(grid[0]) if grid else 0),
            grid=grid, already_tagged=already_tagged,
        ))
    return result


def find_docx_table_grid(docx_path, table_id: str):
    """Tìm bảng vật lý trong 1 file .docx (thường là bienban.docx) có gắn
    tag report_val()/result/... của ĐÚNG table_id ('tables.<table_id>.' xuất
    hiện trong ô nào đó) -> trả grid (list[list[str]], nguyên văn từng ô)
    của bảng đó, ĐÚNG thứ tự dòng/cột thật trong file. None nếu file không
    tồn tại hoặc không bảng nào gắn tag table_id này (mẫu/bảng chưa gắn tag,
    hoặc mã bảng gõ sai) — dùng để dựng bảng rà soát Bước 2/3 khớp ĐÚNG cấu
    trúc file Word thật (xem gui/report_preview.py::build_wysiwyg_table)."""
    docx_path = Path(docx_path)
    if not docx_path.exists():
        return None
    needle = f"tables.{table_id}."
    for detected in scan_docx_tables(docx_path):
        flat = " ".join(c for row in detected.grid for c in row)
        if needle in flat:
            return detected.grid
    return None


def value_columns_from_grid(grid: list, header_row: int = 0) -> list:
    """Trả danh sách chỉ số cột (0-based, ĐÚNG thứ tự trái→phải thật trong
    docx) có ít nhất 1 ô report_val() ở 1 dòng dữ liệu (khác header_row) —
    đây là các cột "giá trị đo" thật, dùng làm tiêu đề cột ở Bước 2/3 thay
    vì tên chung "Lần N" không liên hệ gì tới file thật. Giả định dòng đầu
    tiên (header_row=0) là dòng tiêu đề — đúng với mọi bảng phần mềm này
    từng đọc qua "Đọc bảng từ Word" (mặc định Dòng tiêu đề = 1)."""
    if not grid:
        return []
    n_cols = max(len(r) for r in grid)
    cols = []
    for c in range(n_cols):
        for r_i, row in enumerate(grid):
            if r_i == header_row:
                continue
            if c < len(row) and "report_val()" in row[c]:
                cols.append(c)
                break
    return cols


COLUMN_ROLE_CHOICES = [
    ("none", "(Không dùng)"),
    ("freq_set", "Tần số / điểm đo thiết lập"),
    ("reference", "Chuẩn dùng để tính"),
    ("limit", "Ngưỡng"),
    ("display_label", "Nhãn hiển thị"),
    ("measured", "★ Giá trị đo (phần mềm tự điền tag report_val() — chọn NHIỀU cột nếu 1 dòng có nhiều lần đo)"),
]

_ROLE_KEYWORDS = [
    # Thứ tự ưu tiên — kiểm tra "measured"/"limit" TRƯỚC "freq_set" vì tiêu
    # đề dạng "Tần số đo được" chứa cả "tần số" lẫn "đo được".
    ("measured", ["đo được", "giá trị đo", "kết quả đo", "kết quả"]),
    ("limit", ["ngưỡng", "giới hạn", "sai số cho phép"]),
    ("reference", ["chuẩn"]),
    ("freq_set", ["tần số", "chu kỳ", "danh định", "thiết lập"]),
    ("display_label", ["stt", "điểm đo", "nhãn", "tên", "vị trí", "ký hiệu"]),
]


def guess_column_role(header_text: str) -> str:
    """Đoán vai trò 1 cột theo TỪ KHOÁ trong chữ tiêu đề — CHỈ để gợi ý điền
    sẵn dropdown, người dùng luôn xác nhận/sửa lại trước khi lưu (giống
    scan_meta_paragraphs, best-effort)."""
    t = (header_text or "").lower()
    for role, keywords in _ROLE_KEYWORDS:
        if any(kw in t for kw in keywords):
            return role
    return "none"


def build_rows_from_grid(grid: list, role_map: dict, header_row_index: int = 0,
                          extra_skip_rows: frozenset = frozenset()) -> list:
    """Dựng list[WizardRowSpec] từ lưới text đã đọc (scan_docx_tables) theo
    role_map {cột 0-based: role}, bỏ qua dòng tiêu đề VÀ extra_skip_rows (vd
    1 dòng tiêu đề phụ lồng bên trong bảng, khác dòng tiêu đề chính — xem
    ImportTableFromWordDialog::"Dòng khác không phải dữ liệu"). Số dòng ==
    số dòng dữ liệu thật trong bảng Word (không cần khách gõ lại)."""
    col_by_role = {role: col for col, role in role_map.items() if role and role != "none"}

    def _cell(row: list, role: str) -> str:
        col = col_by_role.get(role)
        if col is None or col >= len(row):
            return ""
        return row[col].strip()

    rows = []
    for r_i, row in enumerate(grid):
        if r_i == header_row_index or r_i in extra_skip_rows:
            continue
        label = _cell(row, "display_label")
        freq_text = _cell(row, "freq_set")
        ref_text = _cell(row, "reference")
        limit_text = _cell(row, "limit")
        key = label or freq_text or f"dòng {r_i + 1}"
        rows.append(WizardRowSpec(
            key=key,
            freq_set=guess_bare_number(freq_text) if freq_text else None,
            reference=guess_bare_number(ref_text) if ref_text else None,
            limit=limit_text,
            display_label=label,
        ))
    return rows


def _measured_grid_positions(tbl, measured_cols, header_row_index: int) -> set:
    """Quy đổi measured_cols (index vào danh sách ô ĐÃ KHỬ TRÙNG LẶP của
    DÒNG TIÊU ĐỀ) sang TOÀN BỘ VỊ TRÍ LƯỚI thật (0-based) mà các ô đó chiếm
    (kể cả khi CHÍNH ô header đó cũng gộp nhiều cột, vd "lần 5" chiếm 2 cột
    lưới — phải tính đủ CẢ 2 vị trí, không chỉ vị trí bắt đầu) — dùng làm
    "toạ độ chung" để so khớp với cấu trúc ô THẬT của từng dòng dữ liệu, vì
    1 dòng có thể gộp ô KHÁC dòng tiêu đề (vd dòng tổng hợp) nên không thể
    dùng chung 1 index đã khử trùng lặp cho mọi dòng."""
    if isinstance(measured_cols, int):
        measured_cols = [measured_cols]
    header_cells = _visible_row_cells_with_pos(tbl.rows[header_row_index])
    positions = set()
    for c in measured_cols:
        if c < len(header_cells):
            _cell, start, span = header_cells[c]
            positions.update(range(start, start + span))
    return positions


def raw_counts_for_measured_cols(docx_path, table_index: int, measured_cols,
                                  header_row_index: int = 0,
                                  extra_skip_rows: frozenset = frozenset()) -> list:
    """[raw_count_dòng_1, raw_count_dòng_2, ...] — DRY RUN (KHÔNG ghi gì) số
    ô "giá trị đo" mỗi dòng dữ liệu THẬT SỰ có, nếu gắn tag theo measured_cols
    (index đã khử trùng lặp ở dòng tiêu đề). Bỏ qua dòng tiêu đề VÀ
    extra_skip_rows (vd 1 dòng tiêu đề phụ lồng bên trong bảng). Số này CÓ
    THỂ KHÁC NHAU giữa các dòng — 1 dòng tổng hợp (vd "Trung Bình") gộp
    nhiều cột "giá trị đo" ở dòng tiêu đề thành 1 ô rộng duy nhất thì CHỈ có
    1 report_val() (không phải N), tự động khớp đúng cấu trúc ô thật, không
    cần cấu hình gì thêm. Dùng để hiển thị trước cho khách xem VÀ để dựng
    RowDef.raw_count đúng từng dòng khi lưu — kết quả khớp CHÍNH XÁC với
    insert_report_val_tags() bên dưới vì dùng chung logic quy đổi vị trí lưới."""
    doc = Document(str(docx_path))
    tbl = doc.tables[table_index]
    measured_positions = _measured_grid_positions(tbl, measured_cols, header_row_index)
    counts = []
    for r_i, row in enumerate(tbl.rows):
        if r_i == header_row_index or r_i in extra_skip_rows:
            continue
        cells = _visible_row_cells_with_pos(row)
        counts.append(sum(1 for _cell, pos, _span in cells if pos in measured_positions))
    return counts


def insert_report_val_tags(docx_path, table_index: int, measured_cols, table_id: str,
                            header_row_index: int = 0,
                            extra_skip_rows: frozenset = frozenset()) -> int:
    """Gõ ĐÈ text của 1 hoặc nhiều cột (measured_cols, index đã khử trùng
    lặp ở DÒNG TIÊU ĐỀ) trong đúng 1 bảng (table_index) thành tag
    `{{ tables.<table_id>.report_val() }}` — bỏ qua dòng tiêu đề VÀ
    extra_skip_rows (vd 1 dòng tiêu đề phụ lồng bên trong bảng, khác dòng
    tiêu đề chính — giữ nguyên chữ tĩnh sẵn có, KHÔNG bị ghi đè), KHÔNG
    thêm/bớt dòng/cột, KHÔNG đụng ô nào khác ngoài phạm vi đã chọn. So khớp
    theo VỊ TRÍ LƯỚI thật (không phải theo index đã khử trùng lặp của TỪNG
    dòng) nên tự động xử lý đúng cả trường hợp 1 dòng gộp ô KHÁC dòng tiêu
    đề (vd dòng tổng hợp "Trung Bình" gộp 5 cột "giá trị đo" hẹp thành 1 ô
    rộng — ô đó chỉ nhận ĐÚNG 1 tag, không phải 5) — số report_val()/dòng vì
    vậy CÓ THỂ KHÁC NHAU giữa các dòng, tự động khớp cấu trúc thật, không
    cần cấu hình thêm và không cần chặn/báo lỗi gì. Trả về tổng số tag đã gắn."""
    doc = Document(str(docx_path))
    tbl = doc.tables[table_index]
    measured_positions = _measured_grid_positions(tbl, measured_cols, header_row_index)
    tag = "{{ tables.%s.report_val() }}" % table_id

    count = 0
    for r_i, row in enumerate(tbl.rows):
        if r_i == header_row_index or r_i in extra_skip_rows:
            continue
        for cell, pos, _span in _visible_row_cells_with_pos(row):
            if pos in measured_positions:
                cell.text = tag
                count += 1
    doc.save(str(docx_path))
    return count


def measured_cell_flags(docx_path, table_index: int, measured_cols, header_row_index: int = 0) -> dict:
    """{row_index (0-based, MỌI dòng kể cả dòng tiêu đề): set(dedup_col_index)}
    — với MỖI dòng, chỉ số ô đã khử trùng lặp (khớp DetectedDocxTable.grid/
    _visible_row_cells) mà vị trí lưới của nó GIAO với vị trí lưới các cột
    "★ Giá trị đo" đã chọn ở dòng tiêu đề (measured_cols) — dùng để CẢNH BÁO
    trước khi ghi đè nếu 1 ô sắp bị gắn tag đang chứa chữ tĩnh thật (xem
    ImportTableFromWordDialog._continue), phòng trường hợp quên tick "Bỏ
    qua" cho 1 dòng tiêu đề phụ lồng bên trong bảng."""
    doc = Document(str(docx_path))
    tbl = doc.tables[table_index]
    measured_positions = _measured_grid_positions(tbl, measured_cols, header_row_index)
    result = {}
    for r_i, row in enumerate(tbl.rows):
        flags = set()
        for c_i, (_cell, start, span) in enumerate(_visible_row_cells_with_pos(row)):
            if any(p in measured_positions for p in range(start, start + span)):
                flags.add(c_i)
        result[r_i] = flags
    return result


# ---------------------------------------------------------------------------
# Định dạng giá trị — form ĐƠN GIẢN (TableFormDialog) KHÔNG bắt khách chọn 1
# trong 19 format nữa (khách phản hồi: quá kỹ thuật, không hiểu) — tự suy
# thẳng từ Đơn vị giá trị đo (khách đã quen — Hz/mVrms/dBm/s/W), LUÔN dùng
# biến thể "không kèm đơn vị" (đơn vị khách tự gõ chữ tĩnh cạnh tag trong
# Word). Chỉ còn 1 lựa chọn phụ: có hiển thị dạng khoa học (×10ⁿ) hay không
# — dùng cho sai số rất nhỏ. FORMAT_LABELS_ALL (19 lựa chọn) vẫn giữ nguyên
# cho RowAdvancedDialog (mỗi report_val() của 1 dòng "nâng cao" có thể cần
# định dạng khác nhau — vd đo được vs sai số kịch bản tự tính) — đó là màn
# hình khác, dành cho ai chủ động bấm "nâng cao", không phải luồng chính.
# ---------------------------------------------------------------------------

_UNIT_TO_FORMAT_NO_UNIT = {
    "Hz": "hz_measured_no_unit",
    "mVrms": "mv_no_unit",
    "dBm": "dbm_no_unit",
    "s": "period_no_unit",
    "W": "w_no_unit",
}


def resolve_value_format(value_unit: str, scientific: bool) -> str:
    """Suy value_format từ Đơn vị giá trị đo — scientific=True ưu tiên hiển
    thị dạng khoa học (×10ⁿ) bất kể đơn vị gì (dùng cho sai số/độ lệch rất
    nhỏ); ngược lại tra theo đơn vị, đơn vị lạ (khách tự gõ) rơi về
    'generic_no_unit' (số thập phân kiểu Việt Nam, không đơn vị)."""
    if scientific:
        return "sci"
    return _UNIT_TO_FORMAT_NO_UNIT.get(value_unit.strip(), "generic_no_unit")


# ---------------------------------------------------------------------------
# Hằng số hiển thị (nhãn tiếng Việt) dùng chung cho màn hình review — thuần
# dữ liệu, không phải code Qt, để ở đây cho gần các hàm dùng chúng.
# ---------------------------------------------------------------------------

FORMAT_LABELS_ALL = [
    ("freq", "Tần số (vd '10 MHz')"),
    ("freq_no_unit", "Tần số — không kèm đơn vị (vd '10')"),
    ("hz_measured", "Tần số đo được (số lẻ đầy đủ)"),
    ("hz_measured_no_unit", "Tần số đo được — không kèm đơn vị"),
    ("period", "Chu kỳ (s)"),
    ("period_no_unit", "Chu kỳ — không kèm đơn vị"),
    ("mv", "mVrms"),
    ("mv_no_unit", "mVrms — không kèm đơn vị"),
    ("dbm", "dBm"),
    ("dbm_no_unit", "dBm — không kèm đơn vị"),
    ("w", "Công suất (W/mW)"),
    ("w_no_unit", "Công suất — không kèm đơn vị"),
    ("sci", "Khoa học (vd '2,4×10⁻⁷')"),
    ("sci_signed", "Khoa học có dấu ± "),
    ("correction_mw", "Số hiệu chỉnh (mW)"),
    ("correction_mw_no_unit", "Số hiệu chỉnh — không kèm đơn vị (mW)"),
    ("correction_db", "Số hiệu chỉnh (dB)"),
    ("correction_db_no_unit", "Số hiệu chỉnh — không kèm đơn vị (dB)"),
    ("text", "Văn bản (giữ nguyên)"),
]

PASS_RULE_CHOICES = [
    ("relative_error_vs_fixed_limit",
     "So sánh sai số tương đối |đo được − chuẩn| / chuẩn với 1 ngưỡng cố định (vd ± 2,4×10⁻⁷)."),
    ("value_vs_parsed_threshold",
     "So sánh giá trị đo với ngưỡng ghi riêng từng dòng (vd '≤ 15 mVrms')."),
    ("correction_vs_reference",
     "Tính số hiệu chỉnh (chuẩn − đo được), KHÔNG có khái niệm đạt/không đạt — dùng cho văn bản hiệu chuẩn."),
    ("none", "Không áp dụng công thức nào (bảng thuần văn bản)."),
]
