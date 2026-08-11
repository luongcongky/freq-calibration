"""
scripts/build_generic_seed_templates.py
==========================================
CHỈ CHẠY 1 LẦN để dựng 2 mẫu báo cáo mới theo ĐÚNG cơ chế thống nhất đã
chốt với khách hàng, qua nhiều vòng góp ý:

1. Mọi field tĩnh qua namespace DUY NHẤT `header.*` (không trộn `meta.*`).
2. **Biên Bản = bản ghi thô** — CHỈ dùng `tables.<ID>.report_val()` (1 mảng
   duy nhất/bảng, đọc tuần tự trái→phải trên→dưới) cho MỌI giá trị đo được;
   KHÔNG có field nào phần mềm tự tính (không report_error()/report_raw()
   kiểu cũ). Giá trị đã biết trước (tần số thiết lập, ngưỡng cố định) vẫn
   là chữ tĩnh, không phải report_val().
3. **GCN = văn bản tổng hợp/kết luận**, được PHÉP dùng field phần mềm tổng
   hợp lại từ dữ liệu report_val Biên Bản ĐÃ CÓ SẴN — vì khách hàng KHÔNG
   THỂ sửa lại hàng loạt kịch bản đã viết xong để thêm dữ liệu riêng cho
   GCN. 2 kiểu tổng hợp:
   - `tables.<ID>.result` — Đạt/Không đạt (mẫu "đo tần số") — người dùng tự
     chọn tay ở Bước 2, không phải công thức tự động.
   - `tables.<ID>.gcn_avg()`/`gcn_error()`/`gcn_limit()` — trung bình/sai số
     (số hiệu chỉnh)/ngưỡng (mẫu "hiệu chuẩn công suất") — tính lại từ
     CHÍNH report_val Biên Bản đã đẩy (map_table()/pass_rule vẫn hoạt động
     y hệt như trước, không cần đổi kịch bản).
   Xem core/table_engine.py::build_cursor_context để biết chi tiết + lý do.

Nội dung 11 bảng (số điểm đo/công thức/ngưỡng) đọc TRỰC TIẾP từ 11 file
JSON descriptor đã có sẵn trong templates/TEMPLATE_FREQ|TEMPLATE_POWER/
tables/.

Riêng bảng A1 (CNT — sai số bộ dao động thạch anh) trước đây cho phép số
lần đo KHÔNG GIỚI HẠN (raw_count=null, vòng lặp docxtpl động) — cơ chế
report_val() cần biết trước SỐ Ô để dựng bảng tĩnh, nên ấn định 5 lần đo
(khớp tiền lệ NRP2 A2/A3 đã dùng raw_count=5) — lệch nhẹ so với nội dung
gốc, đã báo lại với người yêu cầu.

Bảng A1/A5/A6/A7/A8 (QTKĐ_2461_FREQ): mẫu giấy gốc (QTKĐ 2.461:2018 Phụ lục
A) có thêm cột "Sai số" TÍNH TOÁN cạnh giá trị đo được — theo đúng nguyên
tắc report_val(), phần mềm KHÔNG được tự tính field này. Giải pháp đã chốt
với khách hàng: KỊCH BẢN tự tính sai số rồi đẩy THÊM 1 report_val() nữa/dòng
(ngay sau giá trị đo được) — dùng row_def.measured_count để chỉ định máy
chỉ lấy PHẦN ĐẦU raw_readings vào công thức pass_rule (không trộn giá trị
đo + sai số vào 1 phép trung bình), và row_def.value_format_seq để mỗi
report_val() liên tiếp dùng ĐÚNG định dạng của mình. Xem
core/table_engine.py::apply_pass_rule/build_cursor_context. LƯU Ý: A5-A8
hiện CHƯA có kịch bản thật đẩy đủ 2 report_val()/dòng — kịch bản .json
tương lai PHẢI tự tính sai số (không chỉ đẩy số đo thô như trước).

Cả 3 bảng TEMPLATE_POWER: mẫu giấy gốc (QTHC 2.515:2021 Phụ lục A) có
thêm cột "TB" (trung bình, A2/A3 — A1 có dòng "Trung Bình" riêng) và "Độ
KĐBĐ" (độ không đảm bảo đo mở rộng) — cùng nguyên tắc trên: KỊCH BẢN tự
tính (TB tự tính, Độ KĐBĐ theo ngân sách bất định riêng của kỹ sư hiệu
chuẩn — KHÔNG có công thức suy ra từ raw_readings) rồi đẩy thêm 2
report_val()/dòng (TB rồi Độ KĐBĐ, SAU N lần đo). Độ KĐBĐ dùng thêm
row_def.uncertainty_index để lộ ra GCN qua gcn_limit() (không chỉ
report_val() ở Biên Bản) — xem core/table_engine.py::apply_pass_rule. LƯU
Ý: A1/A3 hiện CHƯA có kịch bản thật; A2 đã có kịch bản (cong_suat.json)
nhưng CHƯA đẩy đủ 7 report_val()/dòng (5 đo + TB + Độ KĐBĐ) theo cấu trúc
mới này — cần cập nhật lại kịch bản.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document

from core.table_descriptor import load_table_descriptor
from core.report_generator import _fmt_freq, _fmt_dbm
from core.report_generator_nrp2 import _power_set_from_key
from core.table_layouts import (
    FREQ_A1_HEADERS, FREQ_SENSITIVITY_HEADERS, FREQ_A8_HEADERS, freq_error_headers,
    power_a1_headers, power_a2_headers, power_a3_headers,
)

FREQ_ID = "TEMPLATE_FREQ"
POWER_ID = "TEMPLATE_POWER"
FREQ_DIR = ROOT / "templates" / FREQ_ID
POWER_DIR = ROOT / "templates" / POWER_ID

FREQ_TABLE_IDS = ["A1", "A2", "A3", "A4", "A5", "A6", "A7", "A8"]
POWER_TABLE_IDS = ["A1", "A2", "A3"]

FREQ_A1_RAW_COUNT = 5   # ấn định — xem docstring đầu file
SAI_SO_FORMAT = "sci"   # định dạng cột "Sai số" tính toán (A1/A5-A8) — khoa
                         # học KHÔNG dấu ± (giá trị đo cụ thể, khác "Sai số
                         # cho phép" là 1 ngưỡng tĩnh dùng sci_signed).
# table_id -> list định dạng các report_val() PHỤ (sau giá trị đo được) mà
# KỊCH BẢN phải tự tính rồi đẩy thêm/dòng, theo đúng thứ tự trong bảng thật:
#   A1: fCi (N lần, đã có sẵn) rồi tới fC (TB, kịch bản tự tính, cùng định
#       dạng tần số) rồi δf (sai số, kịch bản tự tính).
#   A5-A8: giá trị đo (đã có sẵn, raw_count=1) rồi tới sai số (kịch bản tự tính).
FREQ_COMPUTED_ERROR_EXTRA_FORMATS = {
    "A1": ["hz_measured", SAI_SO_FORMAT],
    "A5": [SAI_SO_FORMAT], "A6": [SAI_SO_FORMAT],
    "A7": [SAI_SO_FORMAT], "A8": [SAI_SO_FORMAT],
}

# gcn.param_name/limit_str cũ (CNT90XL) — dùng làm chữ TĨNH trong GCN mới.
FREQ_GCN_ROWS = [
    ("A1", "1.Xác định sai số bộ dao động thạch anh", "± 2,4×10⁻⁷"),
    ("A2", "2.Xác định độ nhạy đầu vào kênh A", "Theo QTKĐ"),
    ("A3", "3.Xác định độ nhạy đầu vào kênh B", "Theo QTKĐ"),
    ("A4", "4.Xác định độ nhạy đầu vào kênh C", "Theo QTKĐ"),
    ("A5", "5.Xác định sai số đo tần số kênh A", "± 2,4×10⁻⁷"),
    ("A6", "6.Xác định sai số đo tần số kênh B", "± 2,4×10⁻⁷"),
    ("A7", "7.Xác định sai số đo tần số kênh C", "± 2,4×10⁻⁷"),
    ("A8", "8.Xác định sai số đo chu kỳ", "± 2,4×10⁻⁷"),
]

# value_format đúng đơn vị mỗi bảng (report_val()/gcn_avg() dùng field này).
FREQ_VALUE_FORMAT = {
    "A1": "hz_measured", "A2": "mv", "A3": "mv", "A4": "dbm",
    "A5": "hz_measured", "A6": "hz_measured", "A7": "hz_measured", "A8": "period",
}
POWER_VALUE_FORMAT = {"A1": "w", "A2": "dbm", "A3": "dbm"}


# ---------------------------------------------------------------------------
# Descriptors — đọc từ JSON đã có, đơn giản hoá lại cho đúng cơ chế mới.
# GIỮ NGUYÊN pass_rule + raw_count gốc của từng bảng — vẫn cần để
# map_table() tính value_measured/error/passed nuôi gcn_avg()/gcn_error()/
# result (xem core/table_engine.py) — KHÔNG liên quan tới việc Biên Bản chỉ
# dùng report_val() (report_val() đọc raw_readings, độc lập với pass_rule).
# ---------------------------------------------------------------------------

def _simplify_descriptor(tables_dir: Path, table_id: str, value_format: str,
                          fix_raw_count: int | None = None,
                          computed_error_extra_formats: list | None = None,
                          uncertainty_last_slot: bool = False):
    path = tables_dir / f"{table_id}.json"
    d = load_table_descriptor(path)
    d.columns = []
    d.merge = []
    d.layout = "repeated_rows"
    d.value_format = value_format
    d.gcn = None
    if fix_raw_count is not None:
        for r in d.rows:
            r.raw_count = fix_raw_count
    if computed_error_extra_formats:
        for r in d.rows:
            # Idempotent: nếu file JSON ĐÃ được patch bởi 1 lần chạy script
            # trước (measured_count có sẵn), dùng lại đúng giá trị đó làm
            # gốc — KHÔNG dùng r.raw_count (đã bị cộng dồn thêm extra
            # formats từ lần chạy trước, chạy lại sẽ cộng dồn tiếp nếu lấy
            # nhầm raw_count làm gốc).
            base = r.measured_count if r.measured_count is not None else r.raw_count
            r.measured_count = base
            r.value_format_seq = [value_format] * base + computed_error_extra_formats
            r.raw_count = base + len(computed_error_extra_formats)
            if uncertainty_last_slot:
                r.uncertainty_index = r.raw_count - 1
    path.write_text(json.dumps(d.to_dict(), ensure_ascii=False, indent=2), encoding="utf-8")
    return d


# TEMPLATE_POWER: mỗi dòng đẩy thêm 2 field kịch bản TỰ TÍNH sau N lần đo
# (TB rồi Độ KĐBĐ — xem docstring đầu file). TB cùng value_format của bảng,
# Độ KĐBĐ dùng định dạng "số hiệu chỉnh" đã có sẵn (_ERROR_FORMAT_BY_VALUE_FORMAT
# trong core/table_engine.py) — hợp lý vì cùng là 1 delta nhỏ quanh giá trị đo.
POWER_UNCERTAINTY_FORMAT = {"w": "correction_mw", "dbm": "correction_db"}
POWER_COMPUTED_EXTRA_FORMATS = {
    tid: [POWER_VALUE_FORMAT[tid], POWER_UNCERTAINTY_FORMAT[POWER_VALUE_FORMAT[tid]]]
    for tid in ("A1", "A2", "A3")
}


def build_descriptors():
    freq = {}
    for tid in FREQ_TABLE_IDS:
        fix = FREQ_A1_RAW_COUNT if tid == "A1" else None
        freq[tid] = _simplify_descriptor(FREQ_DIR / "tables", tid, FREQ_VALUE_FORMAT[tid], fix,
                                         FREQ_COMPUTED_ERROR_EXTRA_FORMATS.get(tid))
    power = {}
    for tid in POWER_TABLE_IDS:
        power[tid] = _simplify_descriptor(POWER_DIR / "tables", tid, POWER_VALUE_FORMAT[tid],
                                          computed_error_extra_formats=POWER_COMPUTED_EXTRA_FORMATS[tid],
                                          uncertainty_last_slot=True)
    return freq, power


# ---------------------------------------------------------------------------
# meta.json (không đổi so với lần dựng trước)
# ---------------------------------------------------------------------------

def build_meta_json():
    FREQ_DIR.mkdir(parents=True, exist_ok=True)
    (FREQ_DIR / "meta.json").write_text(json.dumps({
        "template_id": FREQ_ID,
        "template_name": "QTKĐ 2.461 : 2018 — Đo tần số",
        "kind": "kiem_dinh",
        "dut_models": ["CNT-90XL", "CNT-90"],
        "standard": "QTKĐ 2.461 : 2018",
        "measurement_range": "0,002 Hz đến 27 GHz",
        "dut_manufacturer_default": "Pendulum",
    }, ensure_ascii=False, indent=2), encoding="utf-8")

    POWER_DIR.mkdir(parents=True, exist_ok=True)
    (POWER_DIR / "meta.json").write_text(json.dumps({
        "template_id": POWER_ID,
        "template_name": "QTHC 2.515 : 2021 — Hiệu chuẩn công suất",
        "kind": "hieu_chuan",
        "dut_models": ["NRP2"],
        "standard": "QTHC 2.515 : 2021",
        "measurement_range": ("Dải tần làm việc từ DC đến 110 GHz; "
                               "dải đo công suất từ (-67 đến 45) dBm"),
        "dut_manufacturer_default": "R&S",
    }, ensure_ascii=False, indent=2), encoding="utf-8")


# ---------------------------------------------------------------------------
# Helper docx dùng chung
# ---------------------------------------------------------------------------

def _gov_header_table(doc: Document):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = ("CỤC TIÊU CHUẨN ĐO LƯỜNG-CHẤT LƯỢNG / TRUNG TÂM TIÊU CHUẨN "
                            "ĐO LƯỜNG-CHẤT LƯỢNG 2")
    tbl.cell(0, 1).text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM / Độc lập - Tự do - Hạnh phúc"
    return tbl


def _cell_lines(cell, lines: list):
    lines = [ln for ln in lines if ln != ""]
    if not lines:
        return
    cell.paragraphs[0].text = lines[0]
    for ln in lines[1:]:
        cell.add_paragraph(ln)


def _sig_block_2col(doc: Document, label1, tag1, label2, tag2):
    tbl = doc.add_table(rows=1, cols=2)
    _cell_lines(tbl.cell(0, 0), label1.split("\n") + [tag1])
    _cell_lines(tbl.cell(0, 1), label2.split("\n") + [tag2])
    return tbl


def _merge_consecutive(tbl, col: int, start_row: int, n: int, values: list):
    """Gộp các dòng LIÊN TIẾP có cùng giá trị (đã biết trước, tĩnh) tại cột
    `col` — làm THẲNG trong Word lúc dựng file (giá trị đã biết trước, không
    cần chờ render)."""
    i = 0
    while i < n:
        j = i
        while j + 1 < n and values[j + 1] == values[i]:
            j += 1
        r0, r1 = start_row + i, start_row + j
        cell = tbl.cell(r0, col).merge(tbl.cell(r1, col)) if r1 > r0 else tbl.cell(r0, col)
        cell.text = values[i]
        i = j + 1


def _merge_whole_column(tbl, col: int, start_row: int, n: int, text: str):
    cell = tbl.cell(start_row, col).merge(tbl.cell(start_row + n - 1, col)) if n > 1 else tbl.cell(start_row, col)
    cell.text = text
    return cell


# ---------------------------------------------------------------------------
# TEMPLATE_FREQ — bienban.docx (CHỈ report_val(), không có cột tự tính)
# ---------------------------------------------------------------------------

def build_freq_bienban(descriptors: dict):
    doc = Document()
    _gov_header_table(doc)
    doc.add_paragraph("BIÊN BẢN KIỂM ĐỊNH")
    doc.add_paragraph("Tên phương tiện ĐL-TN: {{ header.name }}")
    doc.add_paragraph("Ký hiệu: {{ header.no }}          Số hiệu: {{ header.serial }}")
    doc.add_paragraph("Nước (hãng) sản xuất: {{ header.country }}")
    doc.add_paragraph("Đặc tính đo lường: Dải tần số đo từ {{ header.Characteristics }}")
    doc.add_paragraph("Phương pháp kiểm định: QTKĐ 2.461 : 2018")
    doc.add_paragraph("Phương tiện kiểm định: {{ header.equipment }}")
    doc.add_paragraph("Điều kiện môi trường:  nhiệt độ {{ header.temperature }}"
                       "                độ ẩm {{ header.humidity }}")
    doc.add_paragraph("Đã tiến hành kiểm định ngày {{ header.today }}")
    doc.add_paragraph("KẾT QUẢ KIỂM ĐỊNH")
    doc.add_paragraph("1 Kiểm tra bên ngoài")
    doc.add_paragraph("2 Kiểm tra kỹ thuật")
    doc.add_paragraph("3 Kiểm tra đo lường")

    sections = [
        ("A1", "3.1 Xác định sai số tần số bộ dao động thạch anh"),
        ("A2", "3.2 Xác định độ nhạy đầu vào kênh A"),
        ("A3", "3.3 Xác định độ nhạy đầu vào kênh B"),
        ("A4", "3.4 Xác định độ nhạy đầu vào kênh C"),
        ("A5", "3.5 Xác định sai số đo tần số kênh A"),
        ("A6", "3.6 Xác định sai số đo tần số kênh B"),
        ("A7", "3.7 Xác định sai số đo tần số kênh C"),
        ("A8", "3.8 Xác định sai số đo chu kỳ"),
    ]
    for tid, heading in sections:
        d = descriptors[tid]
        doc.add_paragraph("{%% if tables.%s.enabled %%}" % tid)
        doc.add_paragraph(heading)
        doc.add_paragraph(f"Bảng {tid} - {d.name}")

        if tid == "A1":
            # measured_count = số lần đo fCi thật (raw_count đã +2 cho fC/δf
            # kịch bản tự tính sẵn — xem FREQ_COMPUTED_ERROR_EXTRA_FORMATS).
            n = d.rows[0].measured_count
            headers = FREQ_A1_HEADERS
            # fC/δf ở 1 DÒNG RIÊNG sau khối fCi (KHÔNG gộp dọc xuyên suốt n
            # dòng fCi) — docxtpl đọc tag theo THỨ TỰ VẬT LÝ trong XML, và ô
            # gộp dọc luôn đọc tag ở DÒNG ĐẦU vùng gộp (đã thực nghiệm xác
            # nhận trước đó); nếu gộp fC/δf xuyên suốt khối fCi, tag của
            # chúng bị đọc XEN vào ngay sau fCi dòng 1 thay vì SAU CÙNG,
            # phá thứ tự report_val() kịch bản đã đẩy (đã tự bắt được lỗi
            # này qua test end-to-end thật, không đoán suông).
            tbl = doc.add_table(rows=2 + n, cols=len(headers))
            for j, h in enumerate(headers):
                tbl.cell(0, j).text = h
            _merge_whole_column(tbl, 0, 1, n + 1, _fmt_freq(d.rows[0].freq_set))
            for i in range(n):
                tbl.cell(1 + i, 1).text = "{{ tables.A1.report_val() }}"
            result_row = 1 + n
            tbl.cell(result_row, 2).text = "{{ tables.A1.report_val() }}"
            tbl.cell(result_row, 3).text = "{{ tables.A1.report_val() }}"
            tbl.cell(result_row, 4).text = "± 2,4×10⁻⁷"
        elif tid in ("A2", "A3", "A4"):
            n = len(d.rows)
            tbl = doc.add_table(rows=1 + n, cols=3)
            for j, h in enumerate(FREQ_SENSITIVITY_HEADERS):
                tbl.cell(0, j).text = h
            for i, r in enumerate(d.rows):
                tbl.cell(1 + i, 0).text = _fmt_freq(r.freq_set)
                tbl.cell(1 + i, 1).text = "{{ tables.%s.report_val() }}" % tid
            _merge_consecutive(tbl, 2, 1, n, [r.limit for r in d.rows])
        elif tid in ("A5", "A6", "A7"):
            ch = {"A5": "A", "A6": "B", "A7": "C"}[tid]
            n = len(d.rows)
            headers = freq_error_headers(ch)
            tbl = doc.add_table(rows=1 + n, cols=len(headers))
            for j, h in enumerate(headers):
                tbl.cell(0, j).text = h
            for i, r in enumerate(d.rows):
                tbl.cell(1 + i, 0).text = _fmt_freq(r.freq_set)
                # 2 report_val() liên tiếp/dòng: đo được rồi tới sai số kịch
                # bản tự tính (giống A1 — xem chú thích ở nhánh "A1" trên).
                tbl.cell(1 + i, 1).text = "{{ tables.%s.report_val() }}" % tid
                tbl.cell(1 + i, 2).text = "{{ tables.%s.report_val() }}" % tid
            _merge_whole_column(tbl, 3, 1, n, "± 2,4×10⁻⁷")
        else:  # A8
            n = len(d.rows)
            headers = FREQ_A8_HEADERS
            tbl = doc.add_table(rows=1 + n, cols=len(headers))
            for j, h in enumerate(headers):
                tbl.cell(0, j).text = h
            for i, r in enumerate(d.rows):
                tbl.cell(1 + i, 0).text = r.display_label or r.key
                tbl.cell(1 + i, 1).text = "{{ tables.A8.report_val() }}"
                tbl.cell(1 + i, 2).text = "{{ tables.A8.report_val() }}"
            _merge_whole_column(tbl, 3, 1, n, "± 2,4×10⁻⁷")

        doc.add_paragraph("{% endif %}")

    doc.add_paragraph("4 Kết luận: {{ header.conclusion }}")
    _sig_block_2col(doc, "Người soát lại", "{{ header.reviewer }}",
                     "Kiểm định viên", "{{ header.inspector }}")

    out = FREQ_DIR / "bienban.docx"
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# TEMPLATE_FREQ — gcnkd.docx (bảng tổng hợp: tables.X.result + {%tr if%})
# ---------------------------------------------------------------------------

def build_freq_gcnkd():
    doc = Document()
    _gov_header_table(doc)
    doc.add_paragraph("{{ header.today }}")
    doc.add_paragraph("GIẤY CHỨNG NHẬN KIỂM ĐỊNH")
    doc.add_paragraph("Số: {{ header.cert_no }}")
    doc.add_paragraph("Tên phương tiện ĐL-TN: {{ header.name }}")
    doc.add_paragraph("Ký hiệu: {{ header.no }}          Số hiệu: {{ header.serial }}")
    doc.add_paragraph("Nước (hãng) sản xuất: {{ header.country }}          "
                       "Năm sản xuất: {{ header.birthday }}")
    doc.add_paragraph("Đơn vị sử dụng: {{ header.company }}")
    doc.add_paragraph("Đặc tính đo lường: Dải tần số đo từ {{ header.Characteristics }}.")
    doc.add_paragraph("KẾT QUẢ KIỂM ĐỊNH")

    n_rows = 1 + len(FREQ_GCN_ROWS) * 3
    tbl = doc.add_table(rows=n_rows, cols=3)
    tbl.cell(0, 0).text = "THAM SỐ KIỂM ĐỊNH"
    tbl.cell(0, 1).text = "KẾT QUẢ ĐO"
    tbl.cell(0, 2).text = "GIÁ TRỊ CHO PHÉP"
    r = 1
    for tid, param_name, limit_str in FREQ_GCN_ROWS:
        tbl.cell(r, 0).text = "{%%tr if tables.%s.enabled %%}" % tid
        r += 1
        tbl.cell(r, 0).text = param_name
        tbl.cell(r, 1).text = "{{ tables.%s.result }}" % tid
        tbl.cell(r, 2).text = limit_str
        r += 1
        tbl.cell(r, 0).text = "{%tr endif %}"
        r += 1

    doc.add_paragraph("Phương pháp kiểm định: QTKĐ 2.461 : 2018")
    doc.add_paragraph("Kết luận: {{ header.conclusion }}")
    doc.add_paragraph("Hiệu lực đến {{ header.expire }}")
    _sig_block_2col(doc, "Người kiểm soát\n(Chữ ký, họ tên)", "",
                     "Kiểm định viên\n(Chữ ký, họ tên)", "")
    sig = doc.tables[-1]
    sig.add_column(sig.columns[0].width)
    _cell_lines(sig.cell(0, 2), ["THỦ TRƯỞNG ĐƠN VỊ", "(Ký tên, đóng dấu)", "{{ header.manager }}"])

    doc.add_paragraph("(*) Với điều kiện tôn trọng các nguyên tắc sử dụng và bảo quản"
                       "                    Tổng số trang:………")

    out = FREQ_DIR / "gcnkd.docx"
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# TEMPLATE_POWER — bienban.docx (CHỈ report_val() — từng lần đo thô)
# ---------------------------------------------------------------------------

POWER_A1_GROUP_SIZE = 5   # mẫu giấy QTHC 2.515:2021 chia 10 lần đo A1 thành
                          # 2 nhóm 5 cột (nhóm 1: lần 1-5, nhóm 2: lần 6-10).


def build_power_bienban(descriptors: dict):
    doc = Document()
    _gov_header_table(doc)
    doc.add_paragraph("Thành phố Hồ Chí Minh, ngày      tháng      năm")
    doc.add_paragraph("BIÊN BẢN HIỆU CHUẨN")
    doc.add_paragraph("Số: {{ header.cert_no }}")
    doc.add_paragraph("Tên phương tiện ĐL-TN: {{ header.name }}")
    doc.add_paragraph("Ký hiệu: {{ header.no }}          Số hiệu: {{ header.serial }}")
    doc.add_paragraph("Nước (Hãng) sản xuất: {{ header.country }}          "
                       "Năm sản xuất: {{ header.birthday }}")
    doc.add_paragraph("Đặc tính đo lường: {{ header.Characteristics }}")
    doc.add_paragraph("Đơn vị sử dụng: {{ header.company }}          "
                       "Phương pháp hiệu chuẩn: QTHC 2.515 : 2021")
    doc.add_paragraph("Điều kiện hiệu chuẩn:")
    doc.add_paragraph("Phương tiện hiệu chuẩn: {{ header.equipment }}")
    doc.add_paragraph("Điều kiện môi trường:")
    doc.add_paragraph("- Nhiệt độ: {{ header.temperature }}          - Độ ẩm: {{ header.humidity }}")
    doc.add_paragraph("Đã tiến hành hiệu chuẩn, ngày {{ header.today }}")
    doc.add_paragraph("KẾT QUẢ KIỂM ĐỊNH")
    doc.add_paragraph("A.1 Kiểm tra bên ngoài")
    doc.add_paragraph("A.2 Kiểm tra kỹ thuật")
    doc.add_paragraph("A.2.1 Kiểm tra khả năng làm việc của sensor")
    doc.add_paragraph("A.2.2 Kiểm tra khả năng tự kiểm tra (Self-test)")
    doc.add_paragraph("A.3 Kiểm tra đo lường")

    sections = [
        ("A1", "A.3.1 Xác định độ chính xác mức công suất tại đầu ra chuẩn"),
        ("A2", "A.3.2 Xác định độ chính xác đo mức công suất tuyệt đối (tại 0 dBm)"),
        ("A3", "A.3.3 Xác định độ chính xác đo công suất với bộ hiệu chuẩn công suất "
               "NRPC50 calibration kit"),
    ]
    for tid, heading in sections:
        d = descriptors[tid]
        doc.add_paragraph("{%% if tables.%s.enabled %%}" % tid)
        doc.add_paragraph(heading)
        doc.add_paragraph(f"Bảng {tid} - {d.name}")

        if tid == "A1":
            # measured_count(10)+2 (TB, Độ KĐBĐ kịch bản tự tính — xem
            # POWER_COMPUTED_EXTRA_FORMATS). Chia measured thành các nhóm
            # POWER_A1_GROUP_SIZE cột, TB/Độ KĐBĐ ở 1 DÒNG RIÊNG cuối bảng
            # (không gộp dọc xuyên khối — cùng lý do đã tránh cho A1 FREQ:
            # tag ô gộp dọc bị đọc XEN vào dòng đầu, phá thứ tự report_val()).
            measured = d.rows[0].measured_count
            group = POWER_A1_GROUP_SIZE
            n_groups = -(-measured // group)   # ceil
            label = d.rows[0].display_label or d.rows[0].key
            headers = power_a1_headers(group)
            tbl = doc.add_table(rows=1 + n_groups + 1, cols=len(headers))
            for j, h in enumerate(headers):
                tbl.cell(0, j).text = h
            for g in range(n_groups):
                row = 1 + g
                tbl.cell(row, 0).text = label
                for k in range(group):
                    tbl.cell(row, 1 + k).text = "{{ tables.A1.report_val() }}"
            tb_row = 1 + n_groups
            tbl.cell(tb_row, 0).text = "Trung Bình"
            tb_cell = tbl.cell(tb_row, 1).merge(tbl.cell(tb_row, group))
            tb_cell.text = "{{ tables.A1.report_val() }}"
            tbl.cell(tb_row, group + 1).text = "{{ tables.A1.report_val() }}"
        elif tid == "A2":
            n = len(d.rows)
            raw_n = d.rows[0].measured_count
            headers = power_a2_headers(raw_n)
            tbl = doc.add_table(rows=1 + n, cols=len(headers))
            for j, h in enumerate(headers):
                tbl.cell(0, j).text = h
            for i, r in enumerate(d.rows):
                tbl.cell(1 + i, 0).text = _fmt_freq(r.freq_set)
                for k in range(raw_n):
                    tbl.cell(1 + i, 1 + k).text = "{{ tables.A2.report_val() }}"
                tbl.cell(1 + i, 1 + raw_n).text = "{{ tables.A2.report_val() }}"       # TB
                tbl.cell(1 + i, 2 + raw_n).text = "{{ tables.A2.report_val() }}"       # Độ KĐBĐ
        else:  # A3
            n = len(d.rows)
            raw_n = d.rows[0].measured_count
            headers = power_a3_headers(raw_n)
            tbl = doc.add_table(rows=1 + n, cols=len(headers))
            for j, h in enumerate(headers):
                tbl.cell(0, j).text = h
            freq_labels = [_fmt_freq(r.freq_set) for r in d.rows]
            for i, r in enumerate(d.rows):
                power = _power_set_from_key(r.key)
                tbl.cell(1 + i, 1).text = _fmt_dbm(power) if power is not None else ""
                for k in range(raw_n):
                    tbl.cell(1 + i, 2 + k).text = "{{ tables.A3.report_val() }}"
                tbl.cell(1 + i, 2 + raw_n).text = "{{ tables.A3.report_val() }}"       # TB
                tbl.cell(1 + i, 3 + raw_n).text = "{{ tables.A3.report_val() }}"       # Độ KĐBĐ
            _merge_consecutive(tbl, 0, 1, n, freq_labels)

        doc.add_paragraph("{% endif %}")

    _sig_block_2col(doc, "Người soát lại", "{{ header.reviewer }}",
                     "Kiểm định viên", "{{ header.inspector }}")

    out = POWER_DIR / "bienban.docx"
    doc.save(str(out))
    return out


# ---------------------------------------------------------------------------
# TEMPLATE_POWER — gcnkd.docx. Dựng lại ĐÚNG mẫu Phụ lục B (QTHC 2.515:2021)
# — mẫu giấy CHỈ có đúng 1 dòng "Kết quả (Results):", KHÔNG có bảng A1/A2/A3
# chi tiết nào (khác Biên Bản) — bảng TB/Số hiệu chỉnh/Độ KĐBĐ kiểu cũ đã bị
# BỎ theo đúng yêu cầu khách hàng khớp hệt ảnh mẫu (không có căn cứ từ tài
# liệu — "Trang: 01/03" ở mẫu gợi ý số liệu chi tiết nằm ở trang 2-3, ngoài
# phạm vi Phụ lục B, phần mềm không tự bịa thêm). "Kết quả (Results):" gắn
# {{ header.conclusion }} — ĐÚNG ô "Kết luận" tự động ở Bước 3 (Đạt/Không
# đạt yêu cầu kỹ thuật đo lường, theo Đạt-Không đạt kiểm định viên tự chọn
# tay ở Bước 2 — không phải công thức pass_rule, vì correction_vs_reference
# vốn không có khái niệm đạt/không đạt) — nếu để trống hoàn toàn, GCN xuất
# ra không có kết luận cuối cùng nào cả (đã bị khách hàng phát hiện).
# ---------------------------------------------------------------------------

def build_power_gcnkd():
    doc = Document()
    doc.add_paragraph(
        "CỤC TIÊU CHUẨN ĐO LƯỜNG–CHẤT LƯỢNG (Department for Standard, Metrology and "
        "Quality) / TRUNG TÂM TIÊU CHUẨN ĐO LƯỜNG–CHẤT LƯỢNG 2 (Standard, Metrology and "
        "Quality Centrer 2) / Địa chỉ (Add): Số 40 Nguyễn Giản Thanh – Phường 15 – "
        "Quận10 - TP. Hồ Chí Minh / Điện thoại (Tel): 08.39700977 – Fax: 08.39700977")
    doc.add_paragraph("GIẤY CHỨNG NHẬN HIỆU CHUẨN")
    doc.add_paragraph("(Calibration Certificate)")
    doc.add_paragraph("Số (N0): {{ header.cert_no }}")
    doc.add_paragraph("Tên phương tiện đo (Object): {{ header.name }}")
    doc.add_paragraph("Kiểu (Type): {{ header.no }}          Số (Serial N0): {{ header.serial }}")
    doc.add_paragraph("Nơi sản xuất (Manufacturer): {{ header.country }}")
    doc.add_paragraph("Đặc trưng kỹ thuật (Technical Specification): {{ header.Characteristics }}")
    doc.add_paragraph("Cơ sở sử dụng (Customer): {{ header.company }}")
    doc.add_paragraph("Phương pháp thực hiện (Method of calibration): QTHC 2.515 : 2021")
    doc.add_paragraph("Điều kiện môi trường: (Environmental Conditions):  nhiệt độ "
                       "{{ header.temperature }}          độ ẩm {{ header.humidity }}")
    doc.add_paragraph("Chuẩn được sử dụng (Standards used): {{ header.equipment }}")
    doc.add_paragraph("Kết quả (Results): {{ header.conclusion }}")
    doc.add_paragraph("Ngày hiệu chuẩn (Date of Calibration): {{ header.cal_date }}")
    doc.add_paragraph("")
    doc.add_paragraph("TP Hồ Chí Minh, ngày      tháng      năm 20")

    _sig_block_2col(doc, "TRƯỞNG PHÒNG THÍ NGHIỆM\n(Head of the Cal. Lab.)", "",
                     "GIÁM ĐỐC\n(Director)", "")

    doc.add_paragraph("")
    footer = doc.add_table(rows=1, cols=2)
    _cell_lines(footer.cell(0, 0), ["Trang: 01/03", "(No of paper)"])
    _cell_lines(footer.cell(0, 1), [
        "Không được sao chép rời khi giấy chứng nhận có nhiều trang nếu không được sự "
        "đồng ý bằng văn bản của Trung tâm Tiêu chuẩn-Đo lường-Chất lượng 2",
        "(This certificate shall not be reproduced except in full, without the writen "
        "approval of Standard, Metrology and Quality Centrer 2"])

    out = POWER_DIR / "gcnkd.docx"
    doc.save(str(out))
    return out


def main():
    freq_descriptors, power_descriptors = build_descriptors()
    build_meta_json()
    build_freq_bienban(freq_descriptors)
    build_freq_gcnkd()
    build_power_bienban(power_descriptors)
    build_power_gcnkd()
    print(f"OK - built {FREQ_ID} and {POWER_ID} at {FREQ_DIR} / {POWER_DIR}")


if __name__ == "__main__":
    main()
